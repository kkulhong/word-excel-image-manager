# integrated_word_excel_manager.py
# 통합 Word-Excel 이미지 관리 프로그램
# Tab 1: 이미지 파일명 관리 (report.py)
# Tab 2: Excel 범위 삽입 (excel_to_word_gui.py)

import sys
import os
import shutil
import threading
import traceback
import glob
import locale
import re
import time
import gc
import tempfile
import logging
from datetime import datetime

# PySide6 (Qt) imports
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QLineEdit, QFileDialog, QTextEdit,
    QGroupBox, QProgressBar, QMessageBox, QTableWidget, QTableWidgetItem,
    QHeaderView, QComboBox, QTabWidget, QCheckBox
)
from PySide6.QtCore import Qt, QThread, Signal
from PySide6.QtGui import QFont

# python-docx imports (for Tab 1)
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsmap
from PIL import Image

# win32com imports (for Tab 2)
import win32com.client as win32
import pythoncom

# openpyxl imports (for Tab 2 config)
from openpyxl import Workbook, load_workbook


# ===================================================================
# CONFIGURATION CONSTANTS
# ===================================================================

# Configuration file settings (from excel_to_word_gui.py)
CONFIG_FILE_NAME = "range_config.xlsx"
CONFIG_FILE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), CONFIG_FILE_NAME)

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ===================================================================
# DEFAULT RANGE CONFIGURATION (from excel_to_word_gui.py)
# ===================================================================

DEFAULT_RANGE_CONFIG = {
    "#1": [
        # Below Worst Case Test Ranges
        {"sheet": "GSM 850", "range": "B27:I31", "marker": "GSM 850", "category": "Below Worst Case Test"},
        {"sheet": "W B5", "range": "B27:I31", "marker": "W B5", "category": "Below Worst Case Test"},
        {"sheet": "5B", "range": "B27:J35", "marker": "5B", "category": "Below Worst Case Test"},
        {"sheet": "B5", "range": "B52:J70", "marker": "B5", "category": "Below Worst Case Test"},
        {"sheet": "B12", "range": "B52:J70", "marker": "B12", "category": "Below Worst Case Test"},
        {"sheet": "B13", "range": "B32:J42", "marker": "B13", "category": "Below Worst Case Test"},
        {"sheet": "B14", "range": "B32:J42", "marker": "B14", "category": "Below Worst Case Test"},
        {"sheet": "B26", "range": "B101:J123", "marker": "B26_Part90", "category": "Below Worst Case Test"},
        {"sheet": "B26", "range": "B125:J147", "marker": "B26_Strd", "category": "Below Worst Case Test"},
        {"sheet": "B26", "range": "B149:J171", "marker": "B26_Part22", "category": "Below Worst Case Test"},
        {"sheet": "B71", "range": "B52:J70", "marker": "B71", "category": "Below Worst Case Test"},

        {"sheet": "n5", "range": "B52:K81", "marker": "n5", "category": "Below Worst Case Test"},
        {"sheet": "n12", "range": "B46:K69", "marker": "n12", "category": "Below Worst Case Test"},
        {"sheet": "n14", "range": "B32:K49", "marker": "n14", "category": "Below Worst Case Test"},
        {"sheet": "n26", "range": "B70:K90", "marker": "N26_Part90", "category": "Below Worst Case Test"},
        {"sheet": "n26", "range": "B92:K121", "marker": "N26_Strd", "category": "Below Worst Case Test"},
        {"sheet": "n26", "range": "B123:K152", "marker": "N26_Part22", "category": "Below Worst Case Test"},
        {"sheet": "n71", "range": "B66:K95", "marker": "n71", "category": "Below Worst Case Test"},

        #Below Power Test Ranges
        {"sheet": "GSM 850", "range": "B33:F60", "marker": "GSM 850_Pwr", "category": "Below Power Test"},
        {"sheet": "W B5", "range": "B33:F75", "marker": "W B5_Pwr", "category": "Below Power Test"},
        {"sheet": "5B", "range": "B37:J71", "marker": "5B_Pwr", "category": "Below Power Test"},
        {"sheet": "B5", "range": "B72:H107", "marker": "B5_Pwr1", "category": "Below Power Test"},
        {"sheet": "B5", "range": "B108:H143", "marker": "B5_Pwr2", "category": "Below Power Test"},
        {"sheet": "B12", "range": "B72:H107", "marker": "B12_Pwr1", "category": "Below Power Test"},
        {"sheet": "B12", "range": "B108:H143", "marker": "B12_Pwr2", "category": "Below Power Test"},
        {"sheet": "B13", "range": "B44:H79", "marker": "B13_Pwr", "category": "Below Power Test"},
        {"sheet": "B14", "range": "B44:H79", "marker": "B14_Pwr", "category": "Below Power Test"},
        {"sheet": "B26", "range": "B173:K226", "marker": "B26_Part90_Pwr1", "category": "Below Power Test"},
        {"sheet": "B26", "range": "B227:K262", "marker": "B26_Part90_Pwr2", "category": "Below Power Test"},
        {"sheet": "B71", "range": "B72:H107", "marker": "B71_Pwr", "category": "Below Power Test"},

        {"sheet": "n5", "range": "B83:I138", "marker": "n5_Pwr1", "category": "Below Power Test"},
        {"sheet": "n5", "range": "B139:I178", "marker": "n5_Pwr2", "category": "Below Power Test"},
        {"sheet": "n12", "range": "B71:I126", "marker": "n12_Pwr1", "category": "Below Power Test"},
        {"sheet": "n12", "range": "B127:I146", "marker": "n12_Pwr2", "category": "Below Power Test"},
        {"sheet": "n14", "range": "B51:I106", "marker": "n14_Pwr", "category": "Below Power Test"},
        {"sheet": "n26", "range": "B154:L210", "marker": "N26_Pwr1", "category": "Below Power Test"},
        {"sheet": "n26", "range": "B211:L252", "marker": "N26_Pwr2", "category": "Below Power Test"},
        {"sheet": "n71", "range": "B97:I152", "marker": "n71_Pwr1", "category": "Below Power Test"},
        {"sheet": "n71", "range": "B153:I192", "marker": "n71_Pwr2", "category": "Below Power Test"},

        #Above Worst Case Test Ranges
        {"sheet": "GSM 1900", "range": "B27:I31", "marker": "GSM 1900", "category": "Above Worst Case Test"},
        {"sheet": "W B2", "range": "B27:I31", "marker": "W B2", "category": "Above Worst Case Test"},
        {"sheet": "W B4", "range": "B27:I31", "marker": "W B4", "category": "Above Worst Case Test"},
        {"sheet": "B2", "range": "B64:J90", "marker": "B2", "category": "Above Worst Case Test"},
        {"sheet": "B4", "range": "B64:J90", "marker": "B4", "category": "Above Worst Case Test"},
        {"sheet": "B7", "range": "B52:J70", "marker": "B7", "category": "Above Worst Case Test"},
        {"sheet": "7C", "range": "B27:J35", "marker": "7C", "category": "Above Worst Case Test"},
        {"sheet": "B25", "range": "B64:J90", "marker": "B25", "category": "Above Worst Case Test"},
        {"sheet": "B30", "range": "B34:J44", "marker": "B30", "category": "Above Worst Case Test"},
        {"sheet": "B41", "range": "B52:J70", "marker": "B41", "category": "Above Worst Case Test"},
        {"sheet": "41C", "range": "B27:J37", "marker": "41C", "category": "Above Worst Case Test"},
        {"sheet": "B48", "range": "B52:J70", "marker": "B48", "category": "Above Worst Case Test"},
        {"sheet": "48C", "range": "B27:J37", "marker": "48C", "category": "Above Worst Case Test"},
        {"sheet": "B66", "range": "B64:J90", "marker": "B66", "category": "Above Worst Case Test"},
        {"sheet": "66B", "range": "B52:J70", "marker": "66B", "category": "Above Worst Case Test"},
        {"sheet": "66C", "range": "B27:J37", "marker": "66C", "category": "Above Worst Case Test"},

        {"sheet": "n7", "range": "B82:K141", "marker": "n7", "category": "Above Worst Case Test"},
        {"sheet": "n12", "range": "B46:K69", "marker": "n12", "category": "Above Worst Case Test"},
        {"sheet": "n14", "range": "B38:K55", "marker": "n14", "category": "Above Worst Case Test"},
        {"sheet": "n25", "range": "B76:K129", "marker": "n25", "category": "Above Worst Case Test"},
        {"sheet": "n30", "range": "B36:K53", "marker": "n30", "category": "Above Worst Case Test"},
        {"sheet": "n41", "range": "B112:K201", "marker": "n41", "category": "Above Worst Case Test"},
        {"sheet": "n41_SRS", "range": "B24:I38", "marker": "n41_SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n41_SRS", "range": "M24:T38", "marker": "n41_SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n41_SRS", "range": "X24:AE38", "marker": "n41_SRS4", "category": "Above Worst Case Test"},
        {"sheet": "n48", "range": "B58:K93", "marker": "n48", "category": "Above Worst Case Test"},
        {"sheet": "n48_SRS", "range": "B24:I31", "marker": "n48_SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n48_SRS", "range": "M24:T31", "marker": "n48_SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n48_SRS", "range": "X24:AE31", "marker": "n48_SRS4", "category": "Above Worst Case Test"},
        {"sheet": "n66", "range": "B82:K144", "marker": "n66", "category": "Above Worst Case Test"},
        {"sheet": "n70", "range": "B42:K65", "marker": "n70", "category": "Above Worst Case Test"},
        {"sheet": "n71", "range": "B66:K95", "marker": "n71", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD", "range": "B92:K169", "marker": "n77 DoD", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD SRS", "range": "B22:I36", "marker": "n77 DoD SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD SRS", "range": "M22:T36", "marker": "n77 DoD SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD SRS", "range": "X22:AE36", "marker": "n77 DoD SRS4", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper", "range": "B100:K177", "marker": "n77 Upper", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper SRS", "range": "B24:I38", "marker": "n77 Upper SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper SRS", "range": "M24:T38", "marker": "n77 Upper SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper SRS", "range": "X24:AE38", "marker": "n77 Upper SRS4", "category": "Above Worst Case Test"},

        #Above Power Test Ranges
        {"sheet": "GSM 1900", "range": "B33:F60", "marker": "GSM 1900_Pwr", "category": "Above Power Test"},
        {"sheet": "W B2", "range": "B33:F75", "marker": "W B2_Pwr", "category": "Above Power Test"},
        {"sheet": "W B4", "range": "B33:F75", "marker": "W B4_Pwr", "category": "Above Power Test"},
        {"sheet": "B2", "range": "B92:H145", "marker": "B2_Pwr1", "category": "Above Power Test"},
        {"sheet": "B2", "range": "B146:H199", "marker": "B2_Pwr2", "category": "Above Power Test"},
        {"sheet": "B4", "range": "B92:H145", "marker": "B4_Pwr1", "category": "Above Power Test"},
        {"sheet": "B4", "range": "B146:H199", "marker": "B4_Pwr2", "category": "Above Power Test"},
        {"sheet": "B7", "range": "B72:H107", "marker": "B7_Pwr1", "category": "Above Power Test"},
        {"sheet": "B7", "range": "B108:H143", "marker": "B7_Pwr2", "category": "Above Power Test"},
        {"sheet": "7C", "range": "B37:J71", "marker": "7C_Pwr", "category": "Above Power Test"},
        {"sheet": "B25", "range": "B92:H145", "marker": "B25_Pwr1", "category": "Above Power Test"},
        {"sheet": "B25", "range": "B146:H199", "marker": "B25_Pwr2", "category": "Above Power Test"},
        {"sheet": "B30", "range": "B46:H81", "marker": "B30_Pwr", "category": "Above Power Test"},
        {"sheet": "B41", "range": "B72:H107", "marker": "B41_Pwr1", "category": "Above Power Test"},
        {"sheet": "B41", "range": "B108:H143", "marker": "B41_Pwr2", "category": "Above Power Test"},
        {"sheet": "41C", "range": "B39:J85", "marker": "41C_Pwr", "category": "Above Power Test"},
        {"sheet": "B48", "range": "B72:H107", "marker": "B48_Pwr1", "category": "Above Power Test"},
        {"sheet": "B48", "range": "B108:H143", "marker": "B48_Pwr2", "category": "Above Power Test"},
        {"sheet": "48C", "range": "B39:J85", "marker": "48C_Pwr", "category": "Above Power Test"},
        {"sheet": "B66", "range": "B92:H145", "marker": "B66_Pwr1", "category": "Above Power Test"},
        {"sheet": "B66", "range": "B146:H199", "marker": "B66_Pwr2", "category": "Above Power Test"},
        {"sheet": "66B", "range": "B37:J71", "marker": "66B_Pwr", "category": "Above Power Test"},
        {"sheet": "66C", "range": "B39:J85", "marker": "66C_Pwr", "category": "Above Power Test"},

        {"sheet": "n7", "range": "B82:K141", "marker": "n7", "category": "Above Power Test"},
        {"sheet": "n12", "range": "B46:K69", "marker": "n12", "category": "Above Power Test"},
        {"sheet": "n14", "range": "B38:K55", "marker": "n14", "category": "Above Power Test"},
        {"sheet": "n25", "range": "B76:K129", "marker": "n25", "category": "Above Power Test"},
        {"sheet": "n26", "range": "B70:K90", "marker": "N26_Part90", "category": "Above Power Test"},
        {"sheet": "n26", "range": "B92:K121", "marker": "N26_Strd", "category": "Above Power Test"},
        {"sheet": "n26", "range": "B123:K152", "marker": "N26_Part22", "category": "Above Power Test"},
        {"sheet": "n30", "range": "B36:K53", "marker": "n30", "category": "Above Power Test"},
        {"sheet": "n41", "range": "B112:K201", "marker": "n41", "category": "Above Power Test"},
        {"sheet": "n41_SRS", "range": "B24:I38", "marker": "n41_SRS2", "category": "Above Power Test"},
        {"sheet": "n41_SRS", "range": "M24:T38", "marker": "n41_SRS3", "category": "Above Power Test"},
        {"sheet": "n41_SRS", "range": "X24:AE38", "marker": "n41_SRS4", "category": "Above Power Test"},
        {"sheet": "n48", "range": "B58:K93", "marker": "n48", "category": "Above Power Test"},
        {"sheet": "n48_SRS", "range": "B24:I31", "marker": "n48_SRS2", "category": "Above Power Test"},
        {"sheet": "n48_SRS", "range": "M24:T31", "marker": "n48_SRS3", "category": "Above Power Test"},
        {"sheet": "n48_SRS", "range": "X24:AE31", "marker": "n48_SRS4", "category": "Above Power Test"},
        {"sheet": "n66", "range": "B82:K144", "marker": "n66", "category": "Above Power Test"},
        {"sheet": "n70", "range": "B42:K65", "marker": "n70", "category": "Above Power Test"},
        {"sheet": "n71", "range": "B66:K95", "marker": "n71", "category": "Above Power Test"},
        {"sheet": "n77 DoD", "range": "B92:K169", "marker": "n77 DoD", "category": "Above Power Test"},
        {"sheet": "n77 DoD SRS", "range": "B22:I36", "marker": "n77 DoD SRS2", "category": "Above Power Test"},
        {"sheet": "n77 DoD SRS", "range": "M22:T36", "marker": "n77 DoD SRS3", "category": "Above Power Test"},
        {"sheet": "n77 DoD SRS", "range": "X22:AE36", "marker": "n77 DoD SRS4", "category": "Above Power Test"},
        {"sheet": "n77 Upper", "range": "B100:K177", "marker": "n77 Upper", "category": "Above Power Test"},
        {"sheet": "n77 Upper SRS", "range": "B24:I38", "marker": "n77 Upper SRS2", "category": "Above Power Test"},
        {"sheet": "n77 Upper SRS", "range": "M24:T38", "marker": "n77 Upper SRS3", "category": "Above Power Test"},
        {"sheet": "n77 Upper SRS", "range": "X24:AE38", "marker": "n77 Upper SRS4", "category": "Above Power Test"},
    ],
    "#2": [
        #Below Worst Case Test Ranges
        {"sheet": "GSM 850", "range": "B30:I31", "marker": "GSM 850", "category": "Below Worst Case Test"},
        {"sheet": "W B5", "range": "B33:F75", "marker": "W B5", "category": "Below Worst Case Test"},
        {"sheet": "5B", "range": "B30:J35", "marker": "5B", "category": "Below Worst Case Test"},
        {"sheet": "B5", "range": "B55:J70", "marker": "B5", "category": "Below Worst Case Test"},
        {"sheet": "B12", "range": "B55:J70", "marker": "B12", "category": "Below Worst Case Test"},
        {"sheet": "B13", "range": "B41:J48", "marker": "B13", "category": "Below Worst Case Test"},
        {"sheet": "B14", "range": "B41:J48", "marker": "B14", "category": "Below Worst Case Test"},
        {"sheet": "B26", "range": "B104:J123", "marker": "B26_Part90", "category": "Below Worst Case Test"},
        {"sheet": "B26", "range": "B128:J147", "marker": "B26_Strd", "category": "Below Worst Case Test"},
        {"sheet": "B26", "range": "B152:J171", "marker": "B26_Part22", "category": "Below Worst Case Test"},
        {"sheet": "B71", "range": "B55:J70", "marker": "B71", "category": "Below Worst Case Test"},
        {"sheet": "n5", "range": "B55:K81", "marker": "n5", "category": "Below Worst Case Test"},
        {"sheet": "n12", "range": "B49:K69", "marker": "n12", "category": "Below Worst Case Test"},
        {"sheet": "n14", "range": "B41:K55", "marker": "n14", "category": "Below Worst Case Test"},
        {"sheet": "n26", "range": "B73:K90", "marker": "N26_Part90", "category": "Below Worst Case Test"},
        {"sheet": "n26", "range": "B95:K121", "marker": "N26_Strd", "category": "Below Worst Case Test"},
        {"sheet": "n26", "range": "B126:K152", "marker": "N26_Part22", "category": "Below Worst Case Test"},
        {"sheet": "n71", "range": "B69:K95", "marker": "n71", "category": "Below Worst Case Test"},

        #Below Power Test Ranges
        {"sheet": "GSM 850", "range": "B33:F60", "marker": "GSM 850_Pwr", "category": "Below Power Test"},
        {"sheet": "W B5", "range": "B33:F75", "marker": "W B5_Pwr", "category": "Below Power Test"},
        {"sheet": "5B", "range": "B37:J71", "marker": "5B_Pwr", "category": "Below Power Test"},
        {"sheet": "B5", "range": "B72:H107", "marker": "B5_Pwr1", "category": "Below Power Test"},
        {"sheet": "B5", "range": "B108:H143", "marker": "B5_Pwr2", "category": "Below Power Test"},
        {"sheet": "B12", "range": "B72:H107", "marker": "B12_Pwr1", "category": "Below Power Test"},
        {"sheet": "B12", "range": "B108:H143", "marker": "B12_Pwr2", "category": "Below Power Test"},
        {"sheet": "B13", "range": "B50:H85", "marker": "B13_Pwr", "category": "Below Power Test"},
        {"sheet": "B14", "range": "B50:H85", "marker": "B14_Pwr", "category": "Below Power Test"},
        {"sheet": "B26", "range": "B173:K226", "marker": "B26_Pwr1", "category": "Below Power Test"},
        {"sheet": "B26", "range": "B227:K262", "marker": "B26_Pwr2", "category": "Below Power Test"},
        {"sheet": "B71", "range": "B72:H107", "marker": "B71_Pwr", "category": "Below Power Test"},

        {"sheet": "n5", "range": "B83:I138", "marker": "n5_Pwr1", "category": "Below Power Test"},
        {"sheet": "n5", "range": "B139:I178", "marker": "n5_Pwr2", "category": "Below Power Test"},
        {"sheet": "n12", "range": "B71:I126", "marker": "n12_Pwr1", "category": "Below Power Test"},
        {"sheet": "n12", "range": "B127:I146", "marker": "n12_Pwr2", "category": "Below Power Test"},
        {"sheet": "n14", "range": "B57:I112", "marker": "n14_Pwr", "category": "Below Power Test"},
        {"sheet": "n26", "range": "B154:L210", "marker": "N26_Pwr1", "category": "Below Power Test"},
        {"sheet": "n26", "range": "B211:L252", "marker": "N26_Pwr2", "category": "Below Power Test"},
        {"sheet": "n71", "range": "B97:I152", "marker": "n71_Pwr1", "category": "Below Power Test"},
        {"sheet": "n71", "range": "B153:I192", "marker": "n71_Pwr2", "category": "Below Power Test"},

        #Above Worst Case Test Ranges
        {"sheet": "GSM 1900", "range": "B30:I31", "marker": "GSM 1900", "category": "Above Worst Case Test"},
        {"sheet": "W B2", "range": "B30:I31", "marker": "W B2", "category": "Above Worst Case Test"},
        {"sheet": "W B4", "range": "B30:I31", "marker": "W B4", "category": "Above Worst Case Test"},
        {"sheet": "B2", "range": "B67:J90", "marker": "B2", "category": "Above Worst Case Test"},
        {"sheet": "B4", "range": "B67:J90", "marker": "B4", "category": "Above Worst Case Test"},
        {"sheet": "B7", "range": "B55:J70", "marker": "B7", "category": "Above Worst Case Test"},
        {"sheet": "7C", "range": "B27:J35", "marker": "7C", "category": "Above Worst Case Test"},
        {"sheet": "B25", "range": "B67:J90", "marker": "B25", "category": "Above Worst Case Test"},
        {"sheet": "B30", "range": "B37:J44", "marker": "B30", "category": "Above Worst Case Test"},
        {"sheet": "B41", "range": "B55:J70", "marker": "B41", "category": "Above Worst Case Test"},
        {"sheet": "41C", "range": "B27:J37", "marker": "41C", "category": "Above Worst Case Test"},
        {"sheet": "B48", "range": "B55:J70", "marker": "B48", "category": "Above Worst Case Test"},
        {"sheet": "48C", "range": "B27:J37", "marker": "48C", "category": "Above Worst Case Test"},
        {"sheet": "B66", "range": "B67:J90", "marker": "B66", "category": "Above Worst Case Test"},
        {"sheet": "66B", "range": "B52:J70", "marker": "66B", "category": "Above Worst Case Test"},
        {"sheet": "66C", "range": "B27:J37", "marker": "66C", "category": "Above Worst Case Test"},

        {"sheet": "n7", "range": "B82:K141", "marker": "n7", "category": "Above Worst Case Test"},
        {"sheet": "n12", "range": "B49:K69", "marker": "n12", "category": "Above Worst Case Test"},
        {"sheet": "n14", "range": "B41:K55", "marker": "n14", "category": "Above Worst Case Test"},
        {"sheet": "n25", "range": "B76:K129", "marker": "n25", "category": "Above Worst Case Test"},
        {"sheet": "n30", "range": "B39:K53", "marker": "n30", "category": "Above Worst Case Test"},
        {"sheet": "n41", "range": "B112:K201", "marker": "n41", "category": "Above Worst Case Test"},
        {"sheet": "n41_SRS", "range": "B24:I38", "marker": "n41_SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n41_SRS", "range": "M24:T38", "marker": "n41_SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n41_SRS", "range": "X24:AE38", "marker": "n41_SRS4", "category": "Above Worst Case Test"},
        {"sheet": "n48", "range": "B58:K93", "marker": "n48", "category": "Above Worst Case Test"},
        {"sheet": "n48_SRS", "range": "B24:I31", "marker": "n48_SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n48_SRS", "range": "M24:T31", "marker": "n48_SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n48_SRS", "range": "X24:AE31", "marker": "n48_SRS4", "category": "Above Worst Case Test"},
        {"sheet": "n66", "range": "B82:K144", "marker": "n66", "category": "Above Worst Case Test"},
        {"sheet": "n70", "range": "B45:K65", "marker": "n70", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD", "range": "B92:K169", "marker": "n77 DoD", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD SRS", "range": "B22:I36", "marker": "n77 DoD SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD SRS", "range": "M22:T36", "marker": "n77 DoD SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n77 DoD SRS", "range": "X22:AE36", "marker": "n77 DoD SRS4", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper", "range": "B100:K177", "marker": "n77 Upper", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper SRS", "range": "B24:I38", "marker": "n77 Upper SRS2", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper SRS", "range": "M24:T38", "marker": "n77 Upper SRS3", "category": "Above Worst Case Test"},
        {"sheet": "n77 Upper SRS", "range": "X24:AE38", "marker": "n77 Upper SRS4", "category": "Above Worst Case Test"},

        #Above Power Test Ranges
        {"sheet": "GSM 850", "range": "B33:F60", "marker": "GSM 850_Pwr", "category": "Above Power Test"},
        {"sheet": "W B5", "range": "B33:F75", "marker": "W B5_Pwr", "category": "Above Power Test"},
        {"sheet": "5B", "range": "B37:J71", "marker": "5B_Pwr", "category": "Above Power Test"},
        {"sheet": "B5", "range": "B72:H107", "marker": "B5_Pwr1", "category": "Above Power Test"},
        {"sheet": "B5", "range": "B108:H143", "marker": "B5_Pwr2", "category": "Above Power Test"},
        {"sheet": "B12", "range": "B72:H107", "marker": "B12_Pwr1", "category": "Above Power Test"},
        {"sheet": "B12", "range": "B108:H143", "marker": "B12_Pwr2", "category": "Above Power Test"},
        {"sheet": "B13", "range": "B50:H85", "marker": "B13_Pwr", "category": "Above Power Test"},
        {"sheet": "B14", "range": "B50:H85", "marker": "B14_Pwr", "category": "Above Power Test"},
        {"sheet": "B26", "range": "B173:K226", "marker": "B26_Part90_Pwr1", "category": "Above Power Test"},
        {"sheet": "B26", "range": "B227:K262", "marker": "B26_Part90_Pwr2", "category": "Above Power Test"},
        {"sheet": "B71", "range": "B72:H107", "marker": "B71_Pwr", "category": "Above Power Test"},

        {"sheet": "n5", "range": "B83:I138", "marker": "n5_Pwr1", "category": "Above Power Test"},
        {"sheet": "n5", "range": "B139:I178", "marker": "n5_Pwr2", "category": "Above Power Test"},
        {"sheet": "n12", "range": "B71:I126", "marker": "n12_Pwr1", "category": "Above Power Test"},
        {"sheet": "n12", "range": "B127:I146", "marker": "n12_Pwr2", "category": "Above Power Test"},
        {"sheet": "n14", "range": "B57:I112", "marker": "n14_Pwr", "category": "Above Power Test"},
        {"sheet": "n26", "range": "B154:L210", "marker": "N26_Pwr1", "category": "Above Power Test"},
        {"sheet": "n26", "range": "B211:L252", "marker": "N26_Pwr2", "category": "Above Power Test"},
        {"sheet": "n71", "range": "B97:I152", "marker": "n71_Pwr1", "category": "Above Power Test"},
        {"sheet": "n71", "range": "B153:I192", "marker": "n71_Pwr2", "category": "Above Power Test"},
    ]
}

# Global variable for range config
RANGE_CONFIG = {}


# ===================================================================
# MAIN APPLICATION WINDOW
# ===================================================================

class IntegratedWordExcelManager(QMainWindow):
    """통합 Word-Excel 이미지 관리 프로그램 메인 윈도우"""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("📊 통합 Word-Excel 이미지 관리 프로그램")
        self.setMinimumSize(1200, 850)
        self.resize(1200, 850)

        # Apply global stylesheet
        self.apply_global_styles()

        # Create central widget with tab widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Main layout
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        central_widget.setLayout(main_layout)

        # Create tab widget
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #dcdde1;
                background: #f5f6fa;
            }
            QTabBar::tab {
                background: #ecf0f1;
                color: #2c3e50;
                padding: 10px 20px;
                font-size: 11pt;
                font-weight: bold;
                border: 1px solid #dcdde1;
                border-bottom: none;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background: #3498db;
                color: white;
            }
            QTabBar::tab:hover {
                background: #5dade2;
                color: white;
            }
        """)

        # Create tabs
        self.tab1 = ImageFilenameManagerTab()
        self.tab2 = ExcelRangeInserterTab(self)

        # Connect Tab 2 status update signal to main window status bar
        self.tab2.status_update.connect(self.update_status_bar)

        self.tab_widget.addTab(self.tab1, "📄 이미지 파일명 관리")
        self.tab_widget.addTab(self.tab2, "📊 Excel 범위 삽입")

        main_layout.addWidget(self.tab_widget)

        # Status bar
        self.statusBar().setStyleSheet("""
            QStatusBar {
                background-color: #3498db;
                color: white;
                font-weight: bold;
                padding: 5px;
            }
        """)
        self.statusBar().showMessage("✅ 준비 완료")

        # Add version info to right side of status bar
        version_label = QLabel("v1.0 | 2024-01")
        version_label.setStyleSheet("color: white; padding-right: 10px;")
        self.statusBar().addPermanentWidget(version_label)

    def update_status_bar(self, message):
        """상태 바 업데이트 (Tab 2에서 호출)"""
        self.statusBar().showMessage(message)

    def apply_global_styles(self):
        """전역 스타일 적용"""
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f6fa;
            }
            QGroupBox {
                font-weight: bold;
                font-size: 10pt;
                border: 1px solid #dcdde1;
                border-radius: 6px;
                margin-top: 10px;
                padding-top: 12px;
                background-color: white;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px;
                color: #2c3e50;
            }
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:disabled {
                background-color: #bdc3c7;
            }
            QLineEdit {
                padding: 6px;
                border: 1px solid #dcdde1;
                border-radius: 4px;
                font-size: 9pt;
                background-color: white;
                color: #2c3e50;
            }
            QLineEdit:focus {
                border: 2px solid #3498db;
            }
            QComboBox {
                padding: 5px;
                border: 1px solid #dcdde1;
                border-radius: 4px;
                background-color: white;
                color: #2c3e50;
            }
            QLabel {
                color: #2c3e50;
            }
            QCheckBox {
                color: #2c3e50;
            }
            QTextEdit {
                color: #2c3e50;
                background-color: white;
            }
        """)


# ===================================================================
# TAB 1: IMAGE FILENAME MANAGER (from report.py)
# ===================================================================

class ImageFilenameManagerTab(QWidget):
    """Tab 1: 이미지 파일명 관리 기능"""

    def __init__(self):
        super().__init__()
        self.selected_folder = ""
        self.selected_word_file = ""
        self.include_subfolders = True
        self.worker = None  # QThread worker reference (prevents garbage collection)
        self.btn1 = None  # Function buttons (stored for enable/disable)
        self.btn2 = None
        self.btn3 = None
        self._silent_mode = False  # Worker 스레드용: True면 self.log() 호출 무시 (Qt 스레드 안전성)

        # Set locale for Korean support
        try:
            locale.setlocale(locale.LC_ALL, 'ko_KR.UTF-8')
        except:
            try:
                locale.setlocale(locale.LC_ALL, 'Korean_Korea.949')
            except:
                pass

        self.setup_ui()

    def setup_ui(self):
        """UI 구성"""
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(15, 15, 15, 15)
        main_layout.setSpacing(10)
        self.setLayout(main_layout)

        # 2-column layout (60% left, 40% right)
        content_layout = QHBoxLayout()
        content_layout.setSpacing(10)

        # ========== LEFT COLUMN (60%) ==========
        left_column = QVBoxLayout()
        left_column.setSpacing(10)

        # Folder selection group
        folder_group = QGroupBox("🗂 폴더 선택")
        folder_layout = QVBoxLayout()
        folder_layout.setSpacing(8)

        path_layout = QHBoxLayout()
        path_layout.addWidget(QLabel("폴더 경로:"))
        self.folder_edit = QLineEdit()
        self.folder_edit.setPlaceholderText("이미지 폴더를 선택하세요")
        path_layout.addWidget(self.folder_edit)
        folder_btn = QPushButton("찾기")
        folder_btn.clicked.connect(self.browse_folder)
        path_layout.addWidget(folder_btn)
        folder_layout.addLayout(path_layout)

        self.subfolder_check = QCheckBox("하위폴더 포함")
        self.subfolder_check.setChecked(True)
        folder_layout.addWidget(self.subfolder_check)

        folder_group.setLayout(folder_layout)
        left_column.addWidget(folder_group)

        # Word file selection group
        word_group = QGroupBox("📄 Word 파일 선택")
        word_layout = QHBoxLayout()
        word_layout.addWidget(QLabel("파일 경로:"))
        self.word_edit = QLineEdit()
        self.word_edit.setPlaceholderText("Word 파일을 선택하세요")
        word_layout.addWidget(self.word_edit)
        word_btn = QPushButton("찾기")
        word_btn.clicked.connect(self.browse_word_file)
        word_layout.addWidget(word_btn)
        word_group.setLayout(word_layout)
        left_column.addWidget(word_group)

        # Function selection group
        function_group = QGroupBox("⚙️ 기능 선택")
        function_layout = QVBoxLayout()
        function_layout.setSpacing(8)

        self.btn1 = QPushButton("1. 파일명 기입")
        self.btn1.setMinimumHeight(40)
        self.btn1.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                font-size: 11pt;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        self.btn1.clicked.connect(self.insert_filenames_to_word)
        function_layout.addWidget(self.btn1)

        self.btn2 = QPushButton("2. 이미지 삽입 (통합)")
        self.btn2.setMinimumHeight(40)
        self.btn2.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                font-size: 11pt;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        self.btn2.clicked.connect(self.insert_images_to_word)
        function_layout.addWidget(self.btn2)

        self.btn3 = QPushButton("3. 2열 테이블 자동 생성")
        self.btn3.setMinimumHeight(40)
        self.btn3.setStyleSheet("""
            QPushButton {
                background-color: #e67e22;
                font-size: 11pt;
            }
            QPushButton:hover {
                background-color: #d35400;
            }
        """)
        self.btn3.clicked.connect(self.create_auto_table_with_filenames)
        function_layout.addWidget(self.btn3)

        function_group.setLayout(function_layout)
        left_column.addWidget(function_group)

        content_layout.addLayout(left_column, 60)

        # ========== RIGHT COLUMN (40%) ==========
        right_column = QVBoxLayout()
        right_column.setSpacing(10)

        # Usage guide
        guide_group = QGroupBox("📖 사용 가이드")
        guide_layout = QVBoxLayout()

        self.guide_text = QTextEdit()
        self.guide_text.setReadOnly(True)
        self.guide_text.setMinimumHeight(400)
        self.guide_text.setStyleSheet("""
            QTextEdit {
                background-color: #2c3e50;
                color: #ecf0f1;
                border: 1px solid #34495e;
                border-radius: 4px;
                padding: 10px;
                font-size: 9pt;
                line-height: 1.6;
            }
        """)
        self.guide_text.setHtml("""
<h3 style='color: #3498db;'>📖 이미지 파일명 관리 사용 방법</h3>

<p><b style='color: #27ae60;'>【1. 파일명 기입】</b></p>
<p style='margin-left: 15px;'>
① 이미지 폴더 선택<br>
② Word 문서 선택<br>
③ "1. 파일명 기입" 버튼 클릭<br>
→ Word 표에 이미지 파일명 자동 기입
</p>

<p><b style='color: #3498db;'>【2. 이미지 삽입 (통합)】</b></p>
<p style='margin-left: 15px;'>
① 이미지 폴더 선택<br>
② Word 문서 선택 (표 포함)<br>
③ "2. 이미지 삽입" 버튼 클릭<br>
→ 파일명과 매칭되는 이미지 자동 삽입<br>
→ BE 테스트 셀 자동 감지 및 처리
</p>

<p><b style='color: #e67e22;'>【3. 2열 테이블 자동 생성】</b></p>
<p style='margin-left: 15px;'>
① 이미지 폴더 선택<br>
② "3. 테이블 자동 생성" 버튼 클릭<br>
→ 새 Word 문서에 2열 테이블 생성
</p>

<p><b style='color: #e74c3c;'>💡 팁:</b></p>
<p style='margin-left: 15px;'>
• 하위폴더 포함 시 모든 서브폴더 검색<br>
• BE 테스트: OFDM/DFT-s 자동 인식<br>
• 원본 파일 안전 보관 (_copy 생성)
</p>
        """)
        guide_layout.addWidget(self.guide_text)
        guide_group.setLayout(guide_layout)
        right_column.addWidget(guide_group)

        content_layout.addLayout(right_column, 40)

        main_layout.addLayout(content_layout)

        # ========== BOTTOM SECTION (Full Width) ==========

        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setMinimumHeight(25)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #3498db;
                border-radius: 5px;
                text-align: center;
                background-color: #ecf0f1;
                font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #27ae60;
                border-radius: 4px;
            }
        """)
        main_layout.addWidget(self.progress_bar)

        # Log display
        log_group = QGroupBox("📝 처리 로그")
        log_layout = QVBoxLayout()
        log_layout.setSpacing(8)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMinimumHeight(200)
        self.log_text.setMaximumHeight(250)
        self.log_text.setFont(QFont("Consolas", 9))
        self.log_text.setStyleSheet("""
            QTextEdit {
                background-color: #2c3e50;
                color: #2ecc71;
                border: 1px solid #dcdde1;
                border-radius: 4px;
                padding: 8px;
                font-family: 'Consolas', 'Courier New', monospace;
            }
        """)
        log_layout.addWidget(self.log_text)

        clear_btn = QPushButton("로그 지우기")
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        clear_btn.clicked.connect(self.log_text.clear)
        log_layout.addWidget(clear_btn)

        log_group.setLayout(log_layout)
        main_layout.addWidget(log_group)

    # ========== HELPER METHODS ==========

    def log(self, message):
        """
        로그 메시지 출력

        주의: Worker 스레드에서 호출 시 Qt 스레드 안전성 위반을 방지하기 위해
        _silent_mode가 True면 로그를 출력하지 않습니다.
        """
        # Worker 스레드에서 호출된 경우 GUI 업데이트 건너뛰기 (Qt 스레드 안전성)
        if self._silent_mode:
            return

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.log_text.append(f"[{timestamp}] {message}")
        self.log_text.verticalScrollBar().setValue(
            self.log_text.verticalScrollBar().maximum()
        )
        QApplication.processEvents()

    def windows_sort_key(self, filename):
        """Windows 탐색기 정렬 방식"""
        def convert_part(text):
            if text.isdigit():
                return (0, int(text))
            else:
                return (1, text.lower())

        parts = re.split(r'(\d+)', filename)
        return [convert_part(part) for part in parts if part]

    def browse_folder(self):
        """폴더 선택 다이얼로그"""
        folder = QFileDialog.getExistingDirectory(self, "이미지 폴더 선택")
        if folder:
            self.selected_folder = folder
            self.folder_edit.setText(folder)
            self.log(f"✓ 폴더 선택: {folder}")

    def browse_word_file(self):
        """Word 파일 선택 다이얼로그"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Word 파일 선택",
            "",
            "Word Files (*.docx)"
        )
        if file_path:
            self.selected_word_file = file_path
            self.word_edit.setText(file_path)
            self.log(f"✓ Word 파일 선택: {os.path.basename(file_path)}")

    # ========== HELPER METHODS - Image File Operations ==========

    def get_image_files(self, folder_path, include_subfolders=True, log_callback=None):
        """
        이미지 파일 수집 (하위폴더 포함 옵션)

        주의: 이 메서드는 워커 스레드에서 호출될 수 있으므로 GUI 요소에 직접 접근하지 않습니다.

        Args:
            folder_path: 검색할 폴더 경로
            include_subfolders: 하위 폴더 포함 여부 (GUI에서 전달받은 값)
            log_callback: 로그 출력 콜백 함수 (워커 스레드에서는 시그널 전달)
        """
        image_extensions = ['*.jpg', '*.jpeg', '*.png', '*.gif', '*.bmp', '*.tiff', '*.webp']
        image_files = []

        search_type = "하위폴더 포함" if include_subfolders else "현재 폴더만"
        if log_callback:
            log_callback(f"{search_type} 이미지 파일 검색 중...")

        for ext in image_extensions:
            # GUI 요소에 직접 접근하지 않고 전달받은 파라미터 사용 (Qt 스레드 안전성)
            if include_subfolders:
                pattern = os.path.join(folder_path, '**', ext)
                files = glob.glob(pattern, recursive=True)
            else:
                pattern = os.path.join(folder_path, ext)
                files = glob.glob(pattern, recursive=False)

            # old, etc 포함 폴더 제외
            for file_path in files:
                if not self.is_in_excluded_folder(file_path):
                    image_files.append(file_path)

        filenames = [os.path.basename(f) for f in image_files]
        filenames.sort(key=self.windows_sort_key)

        if log_callback:
            log_callback(f"=== Windows 탐색기 순서로 정렬 ===")
            for i, filename in enumerate(filenames[:10]):
                log_callback(f"{i+1:2d}. {filename}")
            if len(filenames) > 10:
                log_callback(f"    ... 총 {len(filenames)}개 파일")
            log_callback(f"검색 완료 ({search_type}): 총 {len(filenames)}개의 이미지 파일 발견")

        return filenames, image_files

    def get_png_files(self, start_folder, include_subfolders=True, log_callback=None):
        """
        모든 이미지 파일 수집 (PNG, JPG, JPEG 등 + old/etc 폴더 제외)

        주의: 이 메서드는 워커 스레드에서 호출될 수 있으므로 GUI 요소에 직접 접근하지 않습니다.

        Args:
            start_folder: 검색할 폴더 경로
            include_subfolders: 하위 폴더 포함 여부 (GUI에서 전달받은 값)
            log_callback: 로그 출력 콜백 함수 (워커 스레드에서는 시그널 전달)
        """
        png_files = {}
        start_folder = os.path.abspath(start_folder)

        # 모든 이미지 확장자 검색
        image_extensions = ['*.png', '*.jpg', '*.jpeg', '*.gif', '*.bmp', '*.tiff', '*.webp']

        for ext in image_extensions:
            # GUI 요소에 직접 접근하지 않고 전달받은 파라미터 사용 (Qt 스레드 안전성)
            if include_subfolders:
                search_path = os.path.join(start_folder, '**', ext)
                files = glob.glob(search_path, recursive=True)
            else:
                search_path = os.path.join(start_folder, ext)
                files = glob.glob(search_path, recursive=False)

            for file_path in files:
                if os.path.isfile(file_path):
                    # old, etc 폴더 제외
                    if self.is_in_excluded_folder(file_path):
                        # 로그 콜백이 제공된 경우에만 로그 출력 (워커 스레드용)
                        if log_callback:
                            log_callback(f"  제외 폴더 파일 스킵: {os.path.basename(file_path)}")
                        continue

                    name_without_ext = os.path.splitext(os.path.basename(file_path))[0]
                    png_files[name_without_ext] = file_path

        return png_files

    def is_in_excluded_folder(self, file_path):
        """파일이 제외 대상 폴더(old, etc 포함) 안에 있는지 확인"""
        try:
            normalized_path = os.path.normpath(file_path)
            path_parts = normalized_path.split(os.sep)

            # 제외할 폴더 키워드 목록
            excluded_keywords = ['old', 'etc']

            for part in path_parts:
                part_lower = part.lower()
                for keyword in excluded_keywords:
                    if keyword in part_lower:
                        return True

            return False
        except Exception as e:
            self.log(f"  제외 폴더 확인 중 오류: {str(e)}")
            return False

    def create_copy_path(self, original_path, suffix="_copy"):
        """복사본 경로 생성"""
        path_parts = os.path.splitext(original_path)
        copy_path = f"{path_parts[0]}{suffix}{path_parts[1]}"

        counter = 1
        while os.path.exists(copy_path):
            copy_path = f"{path_parts[0]}{suffix}{counter}{path_parts[1]}"
            counter += 1

        return copy_path

    # ========== HELPER METHODS - BE Test Cell Detection ==========

    def is_filename_line(self, text):
        """파일명 라인인지 판단"""
        return (text.startswith('N') and '_' in text and
                ('MHz' in text or 'QPSK' in text or 'QAM' in text or 'DFT' in text or 'CP' in text))

    def is_description_line(self, text):
        """설명 문구인지 판단"""
        description_keywords = [
            'dft-s', 'ofdm', 'qpsk', 'low', 'high', 'frb', 'chansnel', 'chan',
            'spurious', 'emission', 'block', 'error', 'testmode', 'comparison'
        ]
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in description_keywords)

    def is_be_test_cell(self, cell_text):
        """
        BE 테스트 셀 판단 - 극도로 보수적 접근
        BE 전용 키워드가 있을 때만 True
        """
        if not cell_text:
            return False

        # BE 전용 키워드들 (대소문자 무관)
        be_keywords = ['OFDM', 'DFT-s', 'CP_OFDM', 'DFT-s_OFDM']

        text_upper = cell_text.upper()

        # BE 키워드가 하나라도 있으면 BE 셀
        return any(keyword.upper() in text_upper for keyword in be_keywords)

    def update_description_with_testmode(self, text, matched_testmode):
        """TESTMODE에 따라 설명 문구 업데이트"""
        try:
            ofdm_pattern = r'(.*)OFDM(.*)'
            match = re.search(ofdm_pattern, text, re.IGNORECASE)

            if match:
                after_ofdm = match.group(2)

                if matched_testmode == 'DFT':
                    updated_text = f"DFT-s_OFDM{after_ofdm}"
                elif matched_testmode == 'CP':
                    updated_text = f"CP_OFDM{after_ofdm}"
                else:
                    updated_text = text

                return updated_text
            else:
                return text

        except Exception as e:
            self.log(f"      설명 문구 업데이트 중 오류: {str(e)}")
            return text

    # ========== HELPER METHODS - BE Test Cell Processing ==========

    def process_be_comparison_cell(self, cell, png_files, log_callback=None):
        """
        BE 테스트 셀 처리
        Args:
            cell: Word 문서의 셀 객체
            png_files: 이미지 파일 딕셔너리
            log_callback: 로그 출력 콜백 (Worker 스레드용, None이면 self.log 사용)
        """
        try:
            cell_text = cell.text.strip()
            if not cell_text:
                return 0

            lines = [line.strip() for line in cell_text.split('\n') if line.strip()]

            msg = f"      BE 테스트 셀 분석: {len(lines)}개 줄"
            if log_callback:
                log_callback(msg)
            elif not self._silent_mode:
                self.log(msg)

            # 각 줄 분류
            filename_lines = []
            description_lines = []
            other_lines = []

            for line in lines:
                if self.is_filename_line(line):
                    filename_lines.append(line)
                elif self.is_description_line(line):
                    description_lines.append(line)
                else:
                    other_lines.append(line)

            msg = f"      분류: 파일명 {len(filename_lines)}개, 설명 {len(description_lines)}개"
            if log_callback:
                log_callback(msg)
            elif not self._silent_mode:
                self.log(msg)

            # 매칭되는 파일 찾기
            matched_filename = None
            matched_testmode = None

            for filename_line in filename_lines:
                filename_base = filename_line.replace('.png', '').replace('.jpg', '').replace('.jpeg', '')

                if filename_base in png_files:
                    matched_filename = filename_base

                    # TESTMODE 추출
                    if '_DFT' in filename_base.upper():
                        matched_testmode = 'DFT'
                    elif '_CP' in filename_base.upper():
                        matched_testmode = 'CP'

                    msg = f"      ✅ 매칭: {filename_base} (TESTMODE: {matched_testmode})"
                    if log_callback:
                        log_callback(msg)
                    elif not self._silent_mode:
                        self.log(msg)
                    break

            if not matched_filename:
                msg = f"      ❌ 매칭 실패"
                if log_callback:
                    log_callback(msg)
                elif not self._silent_mode:
                    self.log(msg)
                return 0

            # 셀 내용 재구성 (log_callback 전달)
            self.rebuild_be_cell_content(cell, matched_filename, matched_testmode,
                                         description_lines, other_lines, png_files, log_callback=log_callback)

            return 1

        except Exception as e:
            error_msg = f"      ❌ BE 셀 처리 오류: {str(e)}"
            if log_callback:
                log_callback(error_msg)
            elif not self._silent_mode:
                self.log(error_msg)
            return 0

    def rebuild_be_cell_content(self, cell, matched_filename, matched_testmode,
                                description_lines, other_lines, png_files, log_callback=None):
        """
        BE 테스트 셀 내용 재구성 - 공란 완전 제거
        Args:
            cell: Word 문서의 셀 객체
            matched_filename: 매칭된 파일명
            matched_testmode: 추출된 테스트 모드 (DFT/CP)
            description_lines: 설명 줄 리스트
            other_lines: 기타 줄 리스트
            png_files: 이미지 파일 딕셔너리
            log_callback: 로그 출력 콜백 (Worker 스레드용, None이면 self.log 사용)
        """
        try:
            # 원본 설명 문구 서식 저장
            original_desc_formatting = self.save_description_formatting(cell, description_lines)

            # 셀 내용 완전 삭제
            self.clear_cell_safely(cell)

            # 첫 번째 문단도 제거하고 완전히 새로 시작
            tc = cell._tc
            for p_element in list(tc.findall('.//w:p', namespaces=nsmap)):
                tc.remove(p_element)

            # 이미지 삽입 (새 문단 생성) - log_callback 전달하여 에러 메시지 항상 출력
            image_paragraph = cell.add_paragraph()
            image_run = image_paragraph.add_run()

            if self.insert_image_to_run(image_run, png_files[matched_filename], cell.width, log_callback=log_callback):
                msg = f"        ✅ 이미지 삽입: {os.path.basename(png_files[matched_filename])}"
                if log_callback:
                    log_callback(msg)
                elif not self._silent_mode:
                    self.log(msg)
            else:
                error_msg = f"        ❌ 이미지 삽입 실패"
                if log_callback:
                    log_callback(error_msg)
                elif not self._silent_mode:
                    self.log(error_msg)
                return

            # 이미지 문단의 여백 완전 제거
            pf = image_paragraph.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0

            # 기타 텍스트 추가 (내용이 있을 때만)
            for other_line in other_lines:
                if other_line.strip():  # 빈 줄이 아닐 때만
                    other_paragraph = cell.add_paragraph()
                    other_run = other_paragraph.add_run(other_line)
                    pf = other_paragraph.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = 1.0

            # 설명 문구 추가 (내용이 있을 때만)
            for i, desc_line in enumerate(description_lines):
                if desc_line.strip():  # 빈 줄 건너뛰기
                    desc_paragraph = cell.add_paragraph()

                    if matched_testmode and 'OFDM' in desc_line.upper():
                        updated_desc = self.update_description_with_testmode(desc_line, matched_testmode)

                        if updated_desc != desc_line:
                            msg = f"        🔄 설명 업데이트: {desc_line} → {updated_desc}"
                            if log_callback:
                                log_callback(msg)
                            elif not self._silent_mode:
                                self.log(msg)

                        self.add_text_with_original_formatting(desc_paragraph, updated_desc,
                                                              original_desc_formatting, i)
                    else:
                        self.add_text_with_original_formatting(desc_paragraph, desc_line,
                                                              original_desc_formatting, i)

                    # 설명 문단도 여백 제거
                    pf = desc_paragraph.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = 1.0

            msg = f"        ✅ 매칭 안 된 파일명 및 모든 공란 완전 제거 완료"
            if log_callback:
                log_callback(msg)
            elif not self._silent_mode:
                self.log(msg)

        except Exception as e:
            error_msg = f"        ❌ BE 셀 재구성 오류: {str(e)}"
            if log_callback:
                log_callback(error_msg)
            elif not self._silent_mode:
                self.log(error_msg)

    def save_description_formatting(self, cell, description_lines):
        """설명 문구의 원본 서식 정보 저장"""
        formatting_info = []

        try:
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()

                if text in description_lines:
                    para_formatting = {
                        'alignment': paragraph.alignment,
                        'runs': []
                    }

                    for run in paragraph.runs:
                        run_formatting = {
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font_name': None,
                            'font_size': None,
                            'font_color': None
                        }

                        try:
                            if run.font.name:
                                run_formatting['font_name'] = run.font.name
                        except:
                            pass

                        try:
                            if run.font.size:
                                run_formatting['font_size'] = run.font.size
                        except:
                            pass

                        try:
                            if run.font.color and run.font.color.rgb:
                                run_formatting['font_color'] = run.font.color.rgb
                        except:
                            pass

                        para_formatting['runs'].append(run_formatting)

                    formatting_info.append(para_formatting)

        except Exception as e:
            self.log(f"          서식 저장 오류: {str(e)}")

        return formatting_info

    def add_text_with_original_formatting(self, paragraph, text, original_formatting, desc_index):
        """원본 서식을 유지하면서 텍스트만 변경"""
        try:
            if desc_index < len(original_formatting):
                format_info = original_formatting[desc_index]

                if format_info.get('alignment'):
                    paragraph.alignment = format_info['alignment']

                if format_info.get('runs') and len(format_info['runs']) > 0:
                    first_run_format = format_info['runs'][0]
                    new_run = paragraph.add_run(text)

                    if first_run_format.get('bold') is not None:
                        new_run.bold = first_run_format['bold']
                    if first_run_format.get('italic') is not None:
                        new_run.italic = first_run_format['italic']
                    if first_run_format.get('underline') is not None:
                        new_run.underline = first_run_format['underline']

                    try:
                        if first_run_format.get('font_name'):
                            new_run.font.name = first_run_format['font_name']
                    except:
                        pass

                    try:
                        if first_run_format.get('font_size'):
                            new_run.font.size = first_run_format['font_size']
                    except:
                        pass

                    try:
                        if first_run_format.get('font_color'):
                            new_run.font.color.rgb = first_run_format['font_color']
                    except:
                        pass
                else:
                    paragraph.add_run(text)
            else:
                paragraph.add_run(text)

            self.apply_minimal_formatting(paragraph)

        except Exception as e:
            self.log(f"          서식 적용 오류: {str(e)}")
            paragraph.add_run(text)

    def clear_cell_safely(self, cell):
        """셀 내용을 안전하게 삭제"""
        try:
            while len(cell.paragraphs) > 1:
                try:
                    last_paragraph = cell.paragraphs[-1]
                    p_element = last_paragraph._element
                    p_element.getparent().remove(p_element)
                except:
                    break

            if cell.paragraphs:
                first_paragraph = cell.paragraphs[0]
                for run in first_paragraph.runs:
                    run.text = ""
                first_paragraph.text = ""

                try:
                    pf = first_paragraph.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = 1.0
                except:
                    pass

        except Exception as e:
            self.log(f"          셀 삭제 오류: {str(e)}")

    def apply_minimal_formatting(self, paragraph):
        """공백을 최소화하는 문단 서식 적용"""
        try:
            pf = paragraph.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
        except:
            pass

    # ========== HELPER METHODS - Image Insertion ==========

    def find_matching_image(self, text, png_files):
        """이미지 매칭"""
        if not text.strip():
            return None

        return png_files.get(text)

    def get_paragraph_text(self, paragraph):
        """문단 텍스트 추출"""
        return ''.join(run.text for run in paragraph.runs).strip()

    def insert_image_to_run(self, run, img_path, cell_width, log_callback=None):
        """
        이미지를 Run에 삽입

        Args:
            run: Word 문서의 Run 객체
            img_path: 이미지 파일 경로
            cell_width: 셀 너비
            log_callback: 로그 출력 콜백 (Worker 스레드용, None이면 self.log 사용)
                         Worker 스레드에서는 시그널로 전달하여 _silent_mode 영향 회피
        """
        try:
            max_width = Cm(8)

            if hasattr(cell_width, 'cm') and cell_width.cm:
                max_width = Cm(min(cell_width.cm - 0.5, 8))

            with Image.open(img_path) as img:
                width, height = img.size
                aspect_ratio = height / width if width > 0 else 1
                new_width = max_width
                new_height = new_width * aspect_ratio

            run.add_picture(img_path, width=new_width, height=new_height)

            # 성공 로그 - _silent_mode의 영향을 받지 않음
            success_msg = f"    이미지 삽입 성공: {os.path.basename(img_path)}"
            if log_callback:
                log_callback(success_msg)  # Worker 스레드: 시그널로 전달
            elif not self._silent_mode:
                self.log(success_msg)  # 일반 호출: 기본 log 사용
            return True

        except Exception as e:
            # ★★★ 중요: 에러 로그는 반드시 출력되어야 함 (_silent_mode 무시) ★★★
            error_msg = f"    이미지 삽입 실패: {os.path.basename(img_path)}. 오류: {str(e)}"
            if log_callback:
                log_callback(error_msg)  # Worker 스레드: 시그널로 전달 (항상 표시됨)
            elif not self._silent_mode:
                self.log(error_msg)  # 일반 호출: 기본 log 사용
            return False

    def copy_run_format(self, source_run, target_run):
        """Run 서식 복사"""
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def process_cell(self, cell, png_files, log_callback=None):
        """
        일반 셀 처리 로직
        Args:
            cell: Word 문서의 셀 객체
            png_files: 이미지 파일 리스트
            log_callback: 로그 출력 콜백 (Worker 스레드용, None이면 self.log 사용)
        """
        total_attempts = 0
        successful_matches = 0
        successful_insertions = 0

        paragraphs = list(cell.paragraphs)
        for p_idx, paragraph in enumerate(paragraphs):
            try:
                original_text = self.get_paragraph_text(paragraph)
                if not original_text:
                    continue

                total_attempts += 1
                img_path = self.find_matching_image(original_text, png_files)

                if img_path:
                    successful_matches += 1

                    runs_to_process = list(paragraph.runs)
                    paragraph.clear()

                    run = paragraph.add_run()
                    # log_callback을 전달하여 에러 메시지 항상 출력
                    if self.insert_image_to_run(run, img_path, cell.width, log_callback=log_callback):
                        successful_insertions += 1
                        success_msg = f"    ✅ 이미지 매칭 및 삽입 성공: {original_text}"
                        if log_callback:
                            log_callback(success_msg)
                        elif not self._silent_mode:
                            self.log(success_msg)
                    else:
                        error_msg = f"    ❌ 이미지 삽입 실패: {original_text}"
                        if log_callback:
                            log_callback(error_msg)
                        elif not self._silent_mode:
                            self.log(error_msg)
                        for r in runs_to_process:
                            new_run = paragraph.add_run(r.text)
                            self.copy_run_format(r, new_run)

            except Exception as e:
                error_msg = f"    단락 처리 중 오류 발생: {str(e)}"
                if log_callback:
                    log_callback(error_msg)
                elif not self._silent_mode:
                    self.log(error_msg)

        return total_attempts, successful_matches, successful_insertions

    # ========== FEATURE IMPLEMENTATIONS ==========

    def insert_filenames_to_word(self):
        """기능 1: 파일명 기입"""
        if not self.selected_folder or not self.selected_word_file:
            QMessageBox.critical(self, "오류", "폴더와 Word 파일을 모두 선택해주세요.")
            return

        # Check if worker is already running
        if self.worker and self.worker.isRunning():
            QMessageBox.warning(self, "경고", "작업이 이미 진행 중입니다. 완료될 때까지 기다려주세요.")
            return

        # Disable all function buttons during processing
        self.btn1.setEnabled(False)
        self.btn2.setEnabled(False)
        self.btn3.setEnabled(False)

        # Show progress bar
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        # Create and start worker thread (stored as instance variable to prevent garbage collection)
        self.worker = FilenameInsertWorker(
            self.selected_folder,
            self.selected_word_file,
            self.subfolder_check.isChecked(),
            self
        )
        self.worker.progress_update.connect(self.on_progress_update)
        self.worker.log_update.connect(self.log)
        self.worker.finished.connect(self.on_task_finished)
        self.worker.error.connect(self.on_task_error)
        self.worker.start()

    def insert_images_to_word(self):
        """기능 2: 이미지 삽입 (통합)"""
        if not self.selected_folder or not self.selected_word_file:
            QMessageBox.critical(self, "오류", "폴더와 Word 파일을 모두 선택해주세요.")
            return

        # Check if worker is already running
        if self.worker and self.worker.isRunning():
            QMessageBox.warning(self, "경고", "작업이 이미 진행 중입니다. 완료될 때까지 기다려주세요.")
            return

        # Disable all function buttons during processing
        self.btn1.setEnabled(False)
        self.btn2.setEnabled(False)
        self.btn3.setEnabled(False)

        # Show progress bar
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        # Create and start worker thread (stored as instance variable to prevent garbage collection)
        self.worker = ImageInsertWorker(
            self.selected_folder,
            self.selected_word_file,
            self.subfolder_check.isChecked(),
            self
        )
        self.worker.progress_update.connect(self.on_progress_update)
        self.worker.log_update.connect(self.log)
        self.worker.finished.connect(self.on_task_finished)
        self.worker.error.connect(self.on_task_error)
        self.worker.start()

    def create_auto_table_with_filenames(self):
        """기능 3: 2열 테이블 자동 생성"""
        if not self.selected_folder:
            QMessageBox.critical(self, "오류", "이미지 폴더를 선택해주세요.")
            return

        # Check if worker is already running
        if self.worker and self.worker.isRunning():
            QMessageBox.warning(self, "경고", "작업이 이미 진행 중입니다. 완료될 때까지 기다려주세요.")
            return

        # Disable all function buttons during processing
        self.btn1.setEnabled(False)
        self.btn2.setEnabled(False)
        self.btn3.setEnabled(False)

        # Show progress bar
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        # Create and start worker thread (stored as instance variable to prevent garbage collection)
        self.worker = TableCreationWorker(
            self.selected_folder,
            self.subfolder_check.isChecked(),
            self
        )
        self.worker.progress_update.connect(self.on_progress_update)
        self.worker.log_update.connect(self.log)
        self.worker.finished.connect(self.on_task_finished)
        self.worker.error.connect(self.on_task_error)
        self.worker.start()

    # ========== THREAD CALLBACK METHODS ==========

    def on_progress_update(self, value):
        """진행률 업데이트"""
        self.progress_bar.setValue(int(value))

    def on_task_finished(self, message):
        """작업 완료"""
        self.progress_bar.setVisible(False)
        self.log("✅ 작업 완료!")

        # Re-enable all function buttons
        self.btn1.setEnabled(True)
        self.btn2.setEnabled(True)
        self.btn3.setEnabled(True)

        QMessageBox.information(self, "완료", message)

    def on_task_error(self, error_message):
        """작업 오류"""
        self.progress_bar.setVisible(False)
        self.log(f"❌ 오류 발생: {error_message}")

        # Re-enable all function buttons
        self.btn1.setEnabled(True)
        self.btn2.setEnabled(True)
        self.btn3.setEnabled(True)

        QMessageBox.critical(self, "오류", error_message)


# ===================================================================
# WORKER THREADS FOR TAB 1 (QThread implementations)
# ===================================================================

class FilenameInsertWorker(QThread):
    """파일명 기입 작업 스레드"""
    progress_update = Signal(float)
    log_update = Signal(str)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, folder_path, word_file_path, include_subfolders, parent_tab):
        super().__init__()
        self.folder_path = folder_path
        self.word_file_path = word_file_path
        self.include_subfolders = include_subfolders
        self.parent_tab = parent_tab

    def run(self):
        try:
            # ★★★ Qt 스레드 안전성: Worker 스레드에서는 GUI 접근 금지 ★★★
            self.parent_tab._silent_mode = True

            self.progress_update.emit(0)
            self.log_update.emit("=== 파일명 기입 작업 시작 ===")

            # Create backup copy
            original_path = self.word_file_path
            copy_path = self.parent_tab.create_copy_path(original_path)
            self.log_update.emit(f"복사본 생성: {copy_path}")
            shutil.copy2(original_path, copy_path)

            # Get image files - GUI 요소 대신 전달받은 파라미터 사용
            filenames, _ = self.parent_tab.get_image_files(
                self.folder_path,
                include_subfolders=self.include_subfolders,
                log_callback=self.log_update.emit
            )
            if not filenames:
                self.log_update.emit("이미지 파일이 없습니다.")
                self.error.emit("선택한 폴더에 이미지 파일이 없습니다.")
                return

            # Open Word document
            doc = Document(copy_path)
            if not doc.tables:
                self.log_update.emit("Word 문서에 테이블이 없습니다.")
                self.error.emit("Word 문서에 테이블이 없습니다.")
                return

            table = doc.tables[0]
            self.log_update.emit(f"테이블 발견: {len(table.rows)}행 {len(table.columns)}열")

            filename_index = 0
            self.log_update.emit(f"=== 파일명 기입 시작 (Windows 탐색기 순서) ===")

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if filename_index < len(filenames):
                        existing_text = cell.text.strip()
                        filename_without_ext = os.path.splitext(filenames[filename_index])[0]

                        if existing_text:
                            # Save original formatting
                            original_paragraphs_data = []
                            for p in cell.paragraphs:
                                if not p.text.strip():
                                    continue

                                paragraph_data = {'runs': [], 'alignment': p.alignment}
                                for run in p.runs:
                                    run_data = {
                                        "text": run.text,
                                        "bold": run.bold,
                                        "italic": run.italic,
                                        "underline": run.underline,
                                        "font_name": run.font.name,
                                        "font_size": run.font.size,
                                        "font_color_rgb": run.font.color.rgb if run.font.color else None,
                                    }
                                    paragraph_data['runs'].append(run_data)
                                original_paragraphs_data.append(paragraph_data)

                            # Clear cell and rebuild
                            tc = cell._tc
                            for p_element in tc.findall('.//w:p', namespaces=nsmap):
                                tc.remove(p_element)

                            # Add filename
                            p_filename = cell.add_paragraph(filename_without_ext)
                            p_filename.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pf1 = p_filename.paragraph_format
                            pf1.space_before = Pt(0)
                            pf1.space_after = Pt(0)
                            pf1.line_spacing = 1.0

                            # Restore original content
                            for p_data in original_paragraphs_data:
                                p_recreated = cell.add_paragraph()
                                p_recreated.alignment = p_data['alignment']
                                for run_data in p_data['runs']:
                                    new_run = p_recreated.add_run(run_data['text'])
                                    new_run.bold = run_data['bold']
                                    new_run.italic = run_data['italic']
                                    new_run.underline = run_data['underline']
                                    if run_data['font_name']:
                                        new_run.font.name = run_data['font_name']
                                    if run_data['font_size']:
                                        new_run.font.size = run_data['font_size']
                                    if run_data['font_color_rgb']:
                                        new_run.font.color.rgb = run_data['font_color_rgb']

                                pf_recreated = p_recreated.paragraph_format
                                pf_recreated.space_before = Pt(0)
                                pf_recreated.space_after = Pt(0)
                                pf_recreated.line_spacing = 1.0

                            self.log_update.emit(f"셀 서식 포함 재구성 [{row_idx+1},{col_idx+1}]: {filename_without_ext}")

                        else:
                            # Empty cell - just add filename
                            first_paragraph = cell.paragraphs[0]
                            first_paragraph.text = filename_without_ext
                            first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pf = first_paragraph.paragraph_format
                            pf.space_before = Pt(0)
                            pf.space_after = Pt(0)
                            pf.line_spacing = 1.0
                            self.log_update.emit(f"파일명 기입 [{row_idx+1},{col_idx+1}]: {filename_without_ext}")

                        filename_index += 1
                    else:
                        break
                if filename_index >= len(filenames):
                    break

            # Save document
            self.progress_update.emit(100)
            doc.save(copy_path)
            self.log_update.emit(f"파일명 기입 완료! 저장된 파일: {copy_path}")

            self.finished.emit(
                f"파일명이 성공적으로 기입되었습니다.\n\n"
                f"저장된 파일: {os.path.basename(copy_path)}"
            )

        except Exception as e:
            self.error.emit(f"작업 중 오류가 발생했습니다:\n{str(e)}\n\n{traceback.format_exc()}")
        finally:
            # _silent_mode 복원 (Qt 스레드 안전성 정리 작업)
            self.parent_tab._silent_mode = False


class ImageInsertWorker(QThread):
    """이미지 삽입 작업 스레드"""
    progress_update = Signal(float)
    log_update = Signal(str)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, folder_path, word_file_path, include_subfolders, parent_tab):
        super().__init__()
        self.folder_path = folder_path
        self.word_file_path = word_file_path
        self.include_subfolders = include_subfolders
        self.parent_tab = parent_tab

    def run(self):
        try:
            # ★★★ Qt 스레드 안전성: Worker 스레드에서는 GUI 접근 금지 ★★★
            # _silent_mode를 True로 설정하여 헬퍼 메서드의 self.log() 호출 무시
            self.parent_tab._silent_mode = True

            self.progress_update.emit(0)
            self.log_update.emit("=== 통합 이미지 삽입 작업 시작 ===")

            # Create backup copy
            original_path = self.word_file_path
            copy_path = self.parent_tab.create_copy_path(original_path)
            self.log_update.emit(f"복사본 생성: {copy_path}")
            shutil.copy2(original_path, copy_path)

            # Get PNG files - GUI 요소 대신 전달받은 파라미터 사용
            png_files = self.parent_tab.get_png_files(
                self.folder_path,
                include_subfolders=self.include_subfolders,  # GUI에서 전달받은 값 사용
                log_callback=self.log_update.emit  # 로그는 시그널로 전달
            )
            search_type = "하위폴더 포함" if self.include_subfolders else "현재 폴더만"
            self.log_update.emit(f"이미지 파일 검색 완료 ({search_type}): 총 {len(png_files)}개 발견")
            if png_files:
                self.log_update.emit(f"발견된 이미지 파일 (최대 10개): {list(png_files.keys())[:10]}{'...' if len(png_files) > 10 else ''}")

            # Open Word document
            doc = Document(copy_path)
            if not doc.tables:
                self.log_update.emit("Word 문서에 테이블이 없습니다.")
                self.error.emit("Word 문서에 테이블이 없습니다.")
                return

            total_cells = sum(len(row.cells) for table in doc.tables for row in table.rows)
            processed_cells = 0

            # Statistics
            be_test_cells = 0
            basic_cells = 0
            total_be_images = 0
            total_attempts = 0
            total_matches = 0
            total_insertions = 0

            self.log_update.emit(f"총 {len(doc.tables)}개 테이블, {total_cells}개 셀 처리 시작...")

            for table_idx, table in enumerate(doc.tables):
                self.log_update.emit(f"=== 테이블 {table_idx + 1} 처리 중 ===")

                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        processed_cells += 1
                        try:
                            cell_text = cell.text.strip()

                            # ★★★ 핵심 분기 로직 ★★★
                            if self.parent_tab.is_be_test_cell(cell_text):
                                # BE 테스트 셀 처리 - log_callback 전달하여 에러 메시지 항상 출력
                                be_test_cells += 1
                                self.log_update.emit(f"  셀 [{row_idx+1},{col_idx+1}] - BE 테스트 타입 감지")
                                images_inserted = self.parent_tab.process_be_comparison_cell(
                                    cell, png_files, log_callback=self.log_update.emit
                                )
                                total_be_images += images_inserted
                            else:
                                # 일반 셀 처리 - log_callback 전달하여 에러 메시지 항상 출력
                                if cell_text:
                                    basic_cells += 1
                                    self.log_update.emit(f"  셀 [{row_idx+1},{col_idx+1}] - 일반 타입")
                                    attempts, matches, insertions = self.parent_tab.process_cell(
                                        cell, png_files, log_callback=self.log_update.emit
                                    )
                                    total_attempts += attempts
                                    total_matches += matches
                                    total_insertions += insertions

                        except Exception as e:
                            self.log_update.emit(f"  셀 [{row_idx+1},{col_idx+1}] 처리 중 오류: {str(e)}")

                        # Update progress
                        progress_percent = (processed_cells / total_cells) * 100 if total_cells > 0 else 0
                        self.progress_update.emit(progress_percent)

            # Save document
            self.log_update.emit(f"  진행률: 100.0% ({processed_cells}/{total_cells})")
            self.progress_update.emit(100)
            doc.save(copy_path)

            self.log_update.emit("=== 통합 이미지 삽입 완료 ===")
            self.log_update.emit(f"전체 처리 셀: {processed_cells}개")
            self.log_update.emit(f"BE 테스트 셀: {be_test_cells}개 (삽입 이미지: {total_be_images}개)")
            self.log_update.emit(f"일반 셀: {basic_cells}개 (매칭 시도: {total_attempts}, 성공 매칭: {total_matches}, 성공 삽입: {total_insertions})")
            self.log_update.emit(f"저장된 파일: {copy_path}")

            self.finished.emit(
                f"통합 이미지 삽입이 완료되었습니다!\n\n"
                f"전체 처리 셀: {processed_cells:,}개\n"
                f"├─ BE 테스트 셀: {be_test_cells}개\n"
                f"│  └─ 삽입 이미지: {total_be_images}개\n"
                f"└─ 일반 셀: {basic_cells}개\n"
                f"   ├─ 매칭 시도: {total_attempts}개\n"
                f"   ├─ 성공 매칭: {total_matches}개\n"
                f"   └─ 성공 삽입: {total_insertions}개\n\n"
                f"저장된 파일: {os.path.basename(copy_path)}"
            )

        except Exception as e:
            self.error.emit(f"작업 중 오류가 발생했습니다:\n{str(e)}\n\n{traceback.format_exc()}")
        finally:
            # _silent_mode 복원 (Qt 스레드 안전성 정리 작업)
            self.parent_tab._silent_mode = False


class TableCreationWorker(QThread):
    """테이블 자동 생성 작업 스레드"""
    progress_update = Signal(float)
    log_update = Signal(str)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, folder_path, include_subfolders, parent_tab):
        super().__init__()
        self.folder_path = folder_path
        self.include_subfolders = include_subfolders
        self.parent_tab = parent_tab

    def run(self):
        try:
            # ★★★ Qt 스레드 안전성: Worker 스레드에서는 GUI 접근 금지 ★★★
            self.parent_tab._silent_mode = True

            self.progress_update.emit(0)
            self.log_update.emit("=== 2열 테이블 자동 생성 시작 ===")

            # 1. 이미지 파일 수집 - GUI 요소 대신 전달받은 파라미터 사용
            self.log_update.emit("1. 이미지 파일 수집 중...")
            filenames, image_files = self.parent_tab.get_image_files(
                self.folder_path,
                include_subfolders=self.include_subfolders,
                log_callback=self.log_update.emit
            )

            if not filenames:
                self.log_update.emit("이미지 파일이 없습니다.")
                self.error.emit("선택한 폴더에 이미지 파일이 없습니다.")
                return

            self.log_update.emit(f"총 {len(filenames)}개의 이미지 파일 발견")
            self.progress_update.emit(20)

            # 2. 테이블 행 수 계산
            num_images = len(filenames)
            num_rows = (num_images + 1) // 2
            self.log_update.emit(f"2. 테이블 생성: 2열 x {num_rows}행")
            self.progress_update.emit(30)

            # 3. 새 Word 문서 생성
            self.log_update.emit("3. 새 Word 문서 생성 중...")
            doc = Document()
            self.progress_update.emit(40)

            # 4. 2열 N행 테이블 추가
            self.log_update.emit(f"4. {num_rows}행 2열 테이블 추가 중...")
            table = doc.add_table(rows=num_rows, cols=2)
            table.style = 'Table Grid'
            self.progress_update.emit(50)

            # 5. 파일명 기입
            self.log_update.emit("5. 파일명 기입 중 (좌→우, 위→아래 순서)...")
            file_index = 0

            for row_idx in range(num_rows):
                for col_idx in range(2):
                    if file_index < num_images:
                        cell = table.rows[row_idx].cells[col_idx]
                        filename_without_ext = os.path.splitext(filenames[file_index])[0]

                        paragraph = cell.paragraphs[0]
                        paragraph.text = filename_without_ext
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        pf = paragraph.paragraph_format
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)
                        pf.line_spacing = 1.0

                        self.log_update.emit(f"  [{row_idx+1},{col_idx+1}]: {filename_without_ext}")
                        file_index += 1
                    else:
                        self.log_update.emit(f"  [{row_idx+1},{col_idx+1}]: (빈 셀)")
                        break

                # Update progress
                progress = 50 + ((row_idx + 1) / num_rows * 40)
                self.progress_update.emit(progress)

            # 6. 파일 저장
            self.log_update.emit("6. 파일 저장 중...")

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"이미지_테이블_생성_{timestamp}.docx"
            output_path = os.path.join(self.folder_path, output_filename)

            counter = 1
            while os.path.exists(output_path):
                output_filename = f"이미지_테이블_생성_{timestamp}_{counter}.docx"
                output_path = os.path.join(self.folder_path, output_filename)
                counter += 1

            doc.save(output_path)
            self.progress_update.emit(100)

            self.log_update.emit("=== 2열 테이블 자동 생성 완료 ===")
            self.log_update.emit(f"총 이미지 개수: {num_images}개")
            self.log_update.emit(f"테이블 크기: {num_rows}행 x 2열")
            self.log_update.emit(f"저장된 파일: {output_path}")

            self.finished.emit(
                f"2열 테이블이 성공적으로 생성되었습니다!\n\n"
                f"이미지 개수: {num_images}개\n"
                f"테이블 크기: {num_rows}행 x 2열\n\n"
                f"저장된 파일:\n{output_filename}"
            )

        except Exception as e:
            self.error.emit(f"작업 중 오류가 발생했습니다:\n{str(e)}\n\n{traceback.format_exc()}")
        finally:
            # _silent_mode 복원 (Qt 스레드 안전성 정리 작업)
            self.parent_tab._silent_mode = False


# ===================================================================
# TAB 2: EXCEL RANGE INSERTER - WORKER THREAD
# ===================================================================

class ExcelRangeProcessorThread(QThread):
    """백그라운드 작업 스레드 - Excel 범위를 Word에 삽입"""

    progress = Signal(str)
    finished = Signal(dict)

    def __init__(self, excel_files, word_files, mappings):
        super().__init__()
        self.excel_files = excel_files  # 엑셀 파일 리스트
        self.word_files = word_files    # 워드 파일 리스트
        self.mappings = mappings
        self.temp_dir = tempfile.mkdtemp()
        self.output_word_files = []     # 생성된 워드 파일 리스트

    def log(self, message):
        """로그 출력"""
        logger.info(message)
        self.progress.emit(message)

    def create_excel_app_with_retry(self, max_retries=3):
        """Excel 애플리케이션을 재시도 로직과 함께 생성"""
        for attempt in range(max_retries):
            try:
                # COM 초기화 (스레드별로 필요)
                pythoncom.CoInitialize()

                # 첫 시도는 EnsureDispatch, 실패 시 Dispatch 사용
                if attempt == 0:
                    excel = win32.gencache.EnsureDispatch('Excel.Application')
                else:
                    # 캐시 문제 시 Dispatch 사용
                    excel = win32.Dispatch('Excel.Application')

                excel.Visible = False
                excel.DisplayAlerts = False
                self.log(f"  ✓ Excel 애플리케이션 생성 성공 (시도 {attempt + 1}/{max_retries})")
                return excel

            except Exception as e:
                self.log(f"  ⚠️ Excel 생성 실패 (시도 {attempt + 1}/{max_retries}): {str(e)}")

                # COM 정리
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

                # 가비지 컬렉션
                gc.collect()

                if attempt < max_retries - 1:
                    time.sleep(1)  # 재시도 전 대기
                else:
                    raise Exception(f"Excel 애플리케이션 생성 실패 (모든 재시도 실패): {str(e)}")

        return None

    def create_word_app_with_retry(self, max_retries=3):
        """Word 애플리케이션을 재시도 로직과 함께 생성"""
        for attempt in range(max_retries):
            try:
                # COM 초기화 (스레드별로 필요)
                pythoncom.CoInitialize()

                # 첫 시도는 EnsureDispatch, 실패 시 Dispatch 사용
                if attempt == 0:
                    word = win32.gencache.EnsureDispatch('Word.Application')
                else:
                    # 캐시 문제 시 Dispatch 사용
                    word = win32.Dispatch('Word.Application')

                word.Visible = False
                return word

            except Exception as e:
                self.log(f"  ⚠️ Word 생성 실패 (시도 {attempt + 1}/{max_retries}): {str(e)}")

                # COM 정리
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

                # 가비지 컬렉션
                gc.collect()

                if attempt < max_retries - 1:
                    time.sleep(1)  # 재시도 전 대기
                else:
                    raise Exception(f"Word 애플리케이션 생성 실패 (모든 재시도 실패): {str(e)}")

        return None

    def cleanup_com_object(self, obj, obj_name=""):
        """COM 객체 안전하게 정리"""
        if obj is not None:
            try:
                if hasattr(obj, 'Quit'):
                    obj.Quit()
                elif hasattr(obj, 'Close'):
                    obj.Close()
                if obj_name:
                    self.log(f"  ✓ {obj_name} 정리 완료")
            except Exception as e:
                self.log(f"  ⚠️ {obj_name} 정리 중 오류 (무시됨): {str(e)}")

    def create_word_copy(self, word_file):
        """Word 파일 복사본 생성"""
        try:
            # 원본 파일명에서 확장자 분리
            base_name = os.path.splitext(word_file)[0]
            ext = os.path.splitext(word_file)[1]

            # 복사본 파일명 생성
            copy_file = f"{base_name}_copy{ext}"

            # 이미 존재하면 번호 추가
            counter = 1
            while os.path.exists(copy_file):
                copy_file = f"{base_name}_copy{counter}{ext}"
                counter += 1

            # 파일 복사
            shutil.copy2(word_file, copy_file)
            self.log(f"✓ Word 복사본 생성: {os.path.basename(copy_file)}")

            return copy_file

        except Exception as e:
            self.log(f"✗ Word 파일 복사 실패: {str(e)}")
            return None

    def extract_suffix(self, filename):
        """파일명에서 접미사 추출 (#1, #2 등)"""
        match = re.search(r'_#(\d+)\.xlsx?$', filename, re.IGNORECASE)
        if match:
            return f"#{match.group(1)}"
        return None

    def copy_range_as_picture(self, wb, sheet_name, range_address):
        """엑셀 범위를 화면에 보이는 대로 그림으로 복사 (클립보드에)"""
        try:
            # 시트 존재 확인
            try:
                sheet = wb.Sheets(sheet_name)
            except:
                # 시트가 없으면 조용히 건너뜀 (로그 출력 안 함)
                return False

            # 시트가 숨김 상태인지 확인
            # xlSheetVisible = -1 (보임)
            # xlSheetHidden = 0 (숨김)
            # xlSheetVeryHidden = 2 (매우 숨김)
            if sheet.Visible != -1:
                # 숨김 시트도 조용히 건너뜀
                return False

            # 범위 선택
            range_obj = sheet.Range(range_address)

            # 화면에 보이는 대로 그림으로 복사 (클립보드에 복사됨)
            # xlScreen=1, xlPicture=-4147
            range_obj.CopyPicture(Appearance=1, Format=-4147)

            # ★★★ 클립보드 복사 완료 대기 (동기화) ★★★
            # Excel COM의 CopyPicture()는 비동기 작업일 수 있으므로
            # 클립보드에 실제로 복사될 시간을 확보
            time.sleep(0.15)  # 클립보드 안정화 시간 (150ms)

            self.log(f"  ✓ 범위 복사 완료 (클립보드)")
            return True

        except Exception as e:
            # 복사 실패는 조용히 처리 (시트 없음/범위 오류 등)
            return False

    def paste_picture_at_marker(self, word_app, marker):
        """Word 마커 위치에 클립보드의 그림 붙여넣기"""
        try:
            # Selection 초기화
            word_app.Selection.HomeKey(Unit=6)  # wdStory

            # Find 설정
            find = word_app.Selection.Find
            find.ClearFormatting()
            find.Text = marker
            find.Forward = True
            find.Wrap = 1  # wdFindContinue

            # ★★★ 마커 찾기 재시도 로직 (안정성 향상) ★★★
            # Word COM 상태에 따라 첫 시도가 실패할 수 있으므로 최대 2회 시도
            max_find_retries = 2
            marker_found = False

            for retry_attempt in range(max_find_retries):
                if find.Execute():
                    marker_found = True
                    break  # 마커 찾기 성공
                else:
                    if retry_attempt < max_find_retries - 1:
                        # 재시도 전 짧은 대기 및 커서 초기화
                        time.sleep(0.05)
                        word_app.Selection.HomeKey(Unit=6)  # wdStory - 커서 처음으로
                        find.ClearFormatting()
                        find.Text = marker
                        self.log(f"  ⚠️ 마커 찾기 재시도 중... ({retry_attempt + 2}/{max_find_retries})")

            # 마커 찾기 결과 확인
            if marker_found:
                # 페이지 설정 정보
                page_setup = word_app.ActiveDocument.PageSetup
                page_height = page_setup.PageHeight
                top_margin = page_setup.TopMargin
                bottom_margin = page_setup.BottomMargin

                # 마커 위치의 세로 위치 저장 (포인트 단위)
                # 3단계 폴백: Selection.Information(6) → Range.Information(6) → 추정값
                vertical_position = None
                try:
                    vertical_position = word_app.Selection.Information(6)
                except:
                    try:
                        vertical_position = word_app.Selection.Range.Information(6)
                    except:
                        vertical_position = top_margin + 100  # 추정값
                        self.log(f"  ⚠️ 세로 위치 감지 실패, 추정값 사용")

                # 본문 영역 경계 계산
                content_start = top_margin  # 본문 시작 (머릿말 아래)
                content_end = page_height - bottom_margin  # 본문 끝 (바닥글 위)

                # 마커 위치가 본문 영역 내에 있는지 확인
                if vertical_position < content_start:
                    self.log(f"  ⚠️ 경고: 마커가 머릿말 영역에 있음")
                elif vertical_position > content_end:
                    self.log(f"  ⚠️ 경고: 마커가 바닥글 영역에 있음")

                # 안전 여유 공간 (14pt = 약 0.5cm)
                SAFETY_MARGIN = 14

                # 마커부터 본문 끝까지의 거리
                distance_to_content_end = content_end - vertical_position

                # 사용 가능한 높이 = (본문 끝 - 마커 위치 - 안전 여유) × 90%
                raw_available_height = distance_to_content_end - SAFETY_MARGIN
                available_height = raw_available_height * 0.90  # 90% 적용

                self.log(f"  📍 사용 가능 높이: {available_height:.1f}pt ({available_height/28.35:.1f}cm) [90% 적용]")

                # 마커 삭제
                word_app.Selection.Text = ""

                # 클립보드의 그림 붙여넣기
                word_app.Selection.Paste()

                # 방금 붙여넣은 그림 찾기
                picture = None

                # 방법 1: InlineShape 확인 (일반적인 경우)
                if word_app.Selection.InlineShapes.Count > 0:
                    picture = word_app.Selection.InlineShapes(1)

                # 방법 2: Range의 InlineShape 확인 (표 안의 경우)
                elif word_app.Selection.Range.InlineShapes.Count > 0:
                    picture = word_app.Selection.Range.InlineShapes(1)

                # 방법 3: 셀을 선택한 후 InlineShape 확인 (표 셀 안)
                else:
                    try:
                        # 커서를 한 칸 뒤로 이동하여 방금 삽입한 이미지 선택
                        word_app.Selection.MoveLeft(Unit=1, Count=1, Extend=1)
                        if word_app.Selection.InlineShapes.Count > 0:
                            picture = word_app.Selection.InlineShapes(1)
                    except:
                        pass

                if picture is not None:
                    # 원본 크기 저장
                    original_width = picture.Width
                    original_height = picture.Height
                    aspect_ratio = original_height / original_width

                    # 기본 너비를 16.5cm로 설정
                    target_width_cm = 16.5
                    target_width_pt = target_width_cm * 28.35  # 467.72pt

                    # 16.5cm 너비에 맞춘 세로 높이 계산
                    calculated_height = target_width_pt * aspect_ratio

                    if calculated_height > available_height:
                        # 페이지를 넘어가는 경우 → 자동 크기 조정
                        target_height = available_height
                        target_width = target_height / aspect_ratio

                        picture.Height = target_height
                        picture.Width = target_width

                        self.log(f"  ✓ 크기 자동 조정: {target_width/28.35:.1f}cm × {target_height/28.35:.1f}cm")
                    else:
                        # 페이지 내 수용 가능 → 16.5cm 기본 크기 유지
                        picture.Width = target_width_pt
                        picture.Height = calculated_height
                        self.log(f"  ✓ 기본 크기 적용: 16.5cm × {calculated_height/28.35:.1f}cm")
                else:
                    # ★★★ 이미지 객체를 찾을 수 없는 경우 - 명확한 에러 처리 ★★★
                    # picture is None이면 클립보드가 비어있거나 붙여넣기 실패
                    error_msg = f"붙여넣기 후 이미지 객체를 찾을 수 없음 (클립보드 비어있음 또는 COM 오류)"
                    self.log(f"  ❌ 삽입 실패 [{marker}]: {error_msg}")
                    return False, error_msg

                self.log(f"  ✓ 그림 삽입 성공: {marker}")
                return True, None
            else:
                error_msg = f"마커를 Word 문서에서 찾을 수 없음"
                self.log(f"  ✗ 삽입 실패 [{marker}]: {error_msg}")
                return False, error_msg

        except Exception as e:
            error_msg = str(e)
            self.log(f"  ✗ 삽입 실패 [{marker}]: {error_msg}")
            import traceback
            self.log(f"  상세: {traceback.format_exc()}")
            return False, error_msg

    def run(self):
        """메인 처리 - 엑셀-워드 다중 파일 처리"""
        result = {
            'success': False,
            'message': '',
            'images_inserted': 0,
            'images_failed': 0,
            'failed_markers': [],
            'output_files': [],
            'elapsed_time': 0
        }

        # 시작 시간 기록
        start_time = time.time()

        try:
            self.log("=" * 60)
            self.log("엑셀-워드 다중 파일 처리 시작")
            self.log("=" * 60)
            self.log(f"엑셀 파일: {len(self.excel_files)}개")
            self.log(f"워드 파일: {len(self.word_files)}개\n")

            # 워드 파일별로 복사본 생성
            word_copy_files = []
            for word_file in self.word_files:
                copy_file = self.create_word_copy(word_file)
                if copy_file:
                    word_copy_files.append(copy_file)
                    result['output_files'].append(copy_file)
                else:
                    self.log(f"✗ 워드 복사 실패: {os.path.basename(word_file)}")

            if not word_copy_files:
                raise Exception("워드 복사본 생성 실패")

            self.log(f"\n✓ {len(word_copy_files)}개 워드 복사본 생성 완료\n")

            # 워드 파일별로 처리 (Word 기준 방식)
            for word_copy_index, word_copy_file in enumerate(word_copy_files, 1):
                self.log("\n" + "=" * 60)
                self.log(f"[{word_copy_index}/{len(word_copy_files)}] 워드 파일 처리")
                self.log("=" * 60)
                self.log(f"파일: {os.path.basename(word_copy_file)}")

                word = None
                doc = None

                try:
                    # Word 애플리케이션 생성 (한 번만)
                    word = self.create_word_app_with_retry(max_retries=3)
                    doc = word.Documents.Open(os.path.abspath(word_copy_file))
                    self.log(f"✓ 워드 파일 열기 완료")

                    # 모든 엑셀 파일 처리
                    for excel_index, excel_file in enumerate(self.excel_files, 1):
                        self.log(f"\n  [{excel_index}/{len(self.excel_files)}] 엑셀 파일: {os.path.basename(excel_file)}")

                        # 엑셀 파일명에서 접미사 추출
                        suffix = self.extract_suffix(os.path.basename(excel_file))

                        if not suffix:
                            self.log(f"  ⚠️ 접미사를 찾을 수 없음 - 건너뜀")
                            continue

                        self.log(f"  ✓ 접미사: {suffix}")

                        # Excel 열기
                        excel = None
                        wb = None

                        try:
                            # Excel 애플리케이션 생성 (재시도 로직 포함)
                            excel = self.create_excel_app_with_retry(max_retries=3)
                            wb = excel.Workbooks.Open(os.path.abspath(excel_file))
                            self.log(f"  ✓ 엑셀 파일 열기 완료")

                            # RANGE_CONFIG에서 설정 가져오기
                            if suffix in RANGE_CONFIG:
                                config_list = RANGE_CONFIG[suffix]
                                self.log(f"  ✓ {suffix} 설정 사용 ({len(config_list)}개 항목)")

                                for config in config_list:
                                    sheet_name = config['sheet']
                                    range_address = config['range']
                                    marker_prefix = config['marker']
                                    marker = f"{marker_prefix}_{suffix}"

                                    self.log(f"    처리 중: [{sheet_name}] {range_address} → {marker}")

                                    # 엑셀 범위를 그림으로 복사 (클립보드)
                                    if self.copy_range_as_picture(wb, sheet_name, range_address):
                                        # Word에 붙여넣기
                                        success, error_msg = self.paste_picture_at_marker(word, marker)
                                        if success:
                                            result['images_inserted'] += 1
                                            # 이미지 삽입 후 짧은 대기 (Word 과부하 방지 및 안정화)
                                            time.sleep(0.05)
                                        else:
                                            result['images_failed'] += 1
                                            result['failed_markers'].append({
                                                'excel_file': os.path.basename(excel_file),
                                                'word_file': os.path.basename(word_copy_file),
                                                'marker': marker,
                                                'sheet': sheet_name,
                                                'range': range_address,
                                                'reason': error_msg or '알 수 없는 오류'
                                            })
                                    else:
                                        result['images_failed'] += 1
                                        result['failed_markers'].append({
                                            'excel_file': os.path.basename(excel_file),
                                            'word_file': os.path.basename(word_copy_file),
                                            'marker': marker,
                                            'sheet': sheet_name,
                                            'range': range_address,
                                            'reason': '엑셀 범위 복사 실패 (숨김 시트 또는 오류)'
                                        })
                            else:
                                # GUI 테이블 사용 (Fallback)
                                self.log(f"  ⚠️ {suffix} 설정이 없음 - GUI 테이블 사용")

                                for mapping in self.mappings:
                                    sheet_name = mapping['sheet']
                                    range_address = mapping['range']
                                    marker_prefix = mapping['marker']
                                    marker = f"{marker_prefix}_{suffix}"

                                    self.log(f"    처리 중: [{sheet_name}] {range_address} → {marker} (GUI)")

                                    if self.copy_range_as_picture(wb, sheet_name, range_address):
                                        success, error_msg = self.paste_picture_at_marker(word, marker)
                                        if success:
                                            result['images_inserted'] += 1
                                            time.sleep(0.05)
                                        else:
                                            result['images_failed'] += 1
                                            result['failed_markers'].append({
                                                'excel_file': os.path.basename(excel_file),
                                                'word_file': os.path.basename(word_copy_file),
                                                'marker': marker,
                                                'sheet': sheet_name,
                                                'range': range_address,
                                                'reason': error_msg or '알 수 없는 오류'
                                            })
                                    else:
                                        result['images_failed'] += 1
                                        result['failed_markers'].append({
                                            'excel_file': os.path.basename(excel_file),
                                            'word_file': os.path.basename(word_copy_file),
                                            'marker': marker,
                                            'sheet': sheet_name,
                                            'range': range_address,
                                            'reason': '엑셀 범위 복사 실패'
                                        })

                        finally:
                            # 엑셀 정리 (각 Excel 파일 처리 후)
                            if wb is not None:
                                try:
                                    wb.Close(SaveChanges=False)
                                except:
                                    pass

                            if excel is not None:
                                try:
                                    excel.Quit()
                                    self.log(f"  ✓ Excel Application 정리 완료")
                                except Exception as e:
                                    self.log(f"  ⚠️ Excel Application 정리 중 오류: {str(e)}")

                            # 참조 제거
                            wb = None
                            excel = None

                                            # 이미지 삽입 후 짧은 대기 (Word 과부하 방지 및 안정화)
                            # Excel 파일 간 짧은 대기
                            gc.collect()
                            time.sleep(0.3)

                    # 모든 엑셀 파일 처리 완료 - 워드 저장
                    if doc is not None:
                        try:
                            # Word가 안정화될 시간 주기
                            time.sleep(0.5)
                            doc.Save()
                            self.log(f"\n✓ 워드 저장 완료: {os.path.basename(word_copy_file)}")
                        except Exception as e:
                            self.log(f"\n✗ 워드 저장 오류: {str(e)}")
                            # 저장 실패해도 계속 진행 (파일은 이미 수정됨)

                finally:
                    # 워드 정리 (각 Word 파일 처리 후)
                    if doc is not None:
                        try:
                            doc.Close(SaveChanges=False)  # 이미 Save() 호출했으므로
                        except:
                            pass

                    if word is not None:
                        try:
                            word.Quit()
                            self.log(f"✓ Word Application 정리 완료")
                        except Exception as e:
                            self.log(f"⚠️ Word Application 정리 중 오류: {str(e)}")

                    # 참조 제거
                    doc = None
                    word = None

                    # COM 정리 및 가비지 컬렉션 (다음 Word 파일 처리 전)
                    gc.collect()
                    try:
                        pythoncom.CoUninitialize()
                    except:
                        pass
                    time.sleep(1.0)  # Word 파일 간 대기 시간

            # 종료 시간 기록 및 경과 시간 계산
            end_time = time.time()
            elapsed_time = end_time - start_time
            result['elapsed_time'] = elapsed_time

            minutes = int(elapsed_time // 60)
            seconds = int(elapsed_time % 60)

            # 최종 결과 출력
            self.log("\n" + "=" * 60)
            self.log("전체 처리 완료")
            self.log("=" * 60)
            self.log(f"처리 시간: {minutes}분 {seconds}초")
            self.log(f"생성된 워드 파일: {len(result['output_files'])}개")
            self.log(f"삽입 성공: {result['images_inserted']}개")
            self.log(f"삽입 실패: {result['images_failed']}개")

            # 생성된 파일 목록
            if result['output_files']:
                self.log("\n생성된 파일 목록:")
                for idx, file in enumerate(result['output_files'], 1):
                    self.log(f"  {idx}. {os.path.basename(file)}")

            # 실패한 마커 상세 정보
            if result['failed_markers']:
                self.log("\n" + "-" * 60)
                self.log("실패한 마커 목록:")
                self.log("-" * 60)
                for idx, failed in enumerate(result['failed_markers'], 1):
                    self.log(f"{idx}. 엑셀: {failed['excel_file']} → 워드: {failed['word_file']}")
                    self.log(f"   마커: {failed['marker']}, 시트: {failed['sheet']}, 범위: {failed['range']}")
                    self.log(f"   실패 이유: {failed['reason']}")

            result['success'] = True
            result['message'] = "처리 완료"

        except Exception as e:
            self.log(f"\n✗ 치명적 오류 발생: {str(e)}")
            import traceback
            self.log(f"상세 오류:\n{traceback.format_exc()}")
            result['message'] = f"오류: {str(e)}"

        finally:
            # 최종 안전 정리
            self.log("\n최종 정리 작업 중...")

            # 최종 COM 정리
            gc.collect()
            try:
                pythoncom.CoUninitialize()
            except:
                pass

            try:
                os.rmdir(self.temp_dir)
            except:
                pass

        self.finished.emit(result)


# ===================================================================
# TAB 2: EXCEL RANGE INSERTER (from excel_to_word_gui.py)
# ===================================================================

class ExcelRangeInserterTab(QWidget):
    """Tab 2: Excel 범위 삽입 기능"""

    # Signal to update main window status bar
    status_update = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.excel_files = []
        self.word_files = []
        self.mappings = []
        self.worker = None

        # Load configuration
        global RANGE_CONFIG
        RANGE_CONFIG = self.load_or_create_config()

        self.setup_ui()

    def setup_ui(self):
        """UI 구성"""
        main_layout = QVBoxLayout()
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(15, 15, 15, 15)
        self.setLayout(main_layout)

        # 2열 레이아웃 생성 (좌측: 설정/파일/테이블, 우측: 로그)
        content_layout = QHBoxLayout()
        content_layout.setSpacing(10)

        # 좌측 컬럼
        left_column = QVBoxLayout()
        left_column.setSpacing(10)

        # 설정 파일 관리
        config_group = QGroupBox("⚙️ 설정 파일 관리")
        config_layout = QVBoxLayout()
        config_layout.setSpacing(8)
        config_group.setLayout(config_layout)

        # 설정 파일 경로 표시
        config_path_layout = QHBoxLayout()
        config_path_layout.addWidget(QLabel("설정 파일:"))
        self.config_file_label = QLabel(CONFIG_FILE_PATH)
        self.config_file_label.setStyleSheet("color: #27ae60; font-size: 8pt;")
        config_path_layout.addWidget(self.config_file_label)
        config_path_layout.addStretch()
        config_layout.addLayout(config_path_layout)

        # 설정 파일 버튼
        config_btn_layout = QHBoxLayout()
        reload_config_btn = QPushButton("🔄 새로고침")
        reload_config_btn.clicked.connect(self.reload_config_file)
        config_btn_layout.addWidget(reload_config_btn)

        save_config_btn = QPushButton("💾 저장")
        save_config_btn.clicked.connect(self.save_current_config)
        config_btn_layout.addWidget(save_config_btn)

        open_config_btn = QPushButton("📂 열기")
        open_config_btn.clicked.connect(self.open_config_file)
        config_btn_layout.addWidget(open_config_btn)

        config_btn_layout.addStretch()
        config_layout.addLayout(config_btn_layout)

        left_column.addWidget(config_group)

        # 접미사 선택
        suffix_layout = QHBoxLayout()
        suffix_layout.addWidget(QLabel("접미사:"))
        self.suffix_combo = QComboBox()
        self.suffix_combo.addItems(sorted(RANGE_CONFIG.keys()))
        self.suffix_combo.currentTextChanged.connect(self.load_config_to_table)
        self.suffix_combo.setMinimumWidth(100)
        suffix_layout.addWidget(self.suffix_combo)
        load_config_btn = QPushButton("📥 불러오기")
        load_config_btn.clicked.connect(self.load_config_to_table)
        suffix_layout.addWidget(load_config_btn)
        suffix_layout.addStretch()
        left_column.addLayout(suffix_layout)

        # 파일 선택
        file_group = QGroupBox("1️⃣ 파일 선택")
        file_layout = QVBoxLayout()
        file_layout.setSpacing(8)
        file_group.setLayout(file_layout)

        # 엑셀 파일 (다중 선택)
        excel_layout = QHBoxLayout()
        excel_layout.addWidget(QLabel("📊 엑셀:"))
        self.excel_edit = QLineEdit()
        self.excel_edit.setPlaceholderText("여러 파일 선택 가능 (예: 파일명_#1.xlsx, 파일명_#2.xlsx)")
        excel_layout.addWidget(self.excel_edit)
        excel_btn = QPushButton("찾기")
        excel_btn.setMinimumWidth(70)
        excel_btn.clicked.connect(self.select_excel_files)
        excel_layout.addWidget(excel_btn)
        file_layout.addLayout(excel_layout)

        # 선택된 파일 목록 표시
        self.selected_files_label = QLabel("선택된 파일: 0개")
        self.selected_files_label.setStyleSheet("color: #27ae60; font-weight: bold;")
        file_layout.addWidget(self.selected_files_label)

        # Word 파일 (다중 선택)
        word_layout = QHBoxLayout()
        word_layout.addWidget(QLabel("📄 워드:"))
        self.word_edit = QLineEdit()
        self.word_edit.setPlaceholderText("여러 파일 선택 가능 (예: GSM850.docx, B14.docx)")
        word_layout.addWidget(self.word_edit)
        word_btn = QPushButton("찾기")
        word_btn.setMinimumWidth(70)
        word_btn.clicked.connect(self.select_word_files)
        word_layout.addWidget(word_btn)
        file_layout.addLayout(word_layout)

        # 선택된 워드 파일 목록 표시
        self.selected_word_files_label = QLabel("선택된 파일: 0개")
        self.selected_word_files_label.setStyleSheet("color: #3498db; font-weight: bold;")
        file_layout.addWidget(self.selected_word_files_label)

        left_column.addWidget(file_group)

        # 매핑 설정
        mapping_group = QGroupBox("2️⃣ 복사 범위 설정")
        mapping_layout = QVBoxLayout()
        mapping_layout.setSpacing(8)
        mapping_group.setLayout(mapping_layout)

        # 설명
        info = QLabel("💡 엑셀의 범위를 Word의 마커에 삽입 | 마커 형식: '마커접두사_#1' (예: 'GSM 850_#1')")
        info.setStyleSheet("color: #7f8c8d; font-size: 8pt; padding: 5px;")
        mapping_layout.addWidget(info)

        # 테이블
        self.mapping_table = QTableWidget()
        self.mapping_table.setColumnCount(4)
        self.mapping_table.setHorizontalHeaderLabels(['시트명', '범위', '마커접두사', '카테고리'])
        self.mapping_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.mapping_table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #dcdde1;
                background-color: white;
                gridline-color: #ecf0f1;
                alternate-background-color: #f8f9fa;
            }
            QHeaderView::section {
                background-color: #3498db;
                color: white;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
        """)
        self.mapping_table.setAlternatingRowColors(True)

        # RANGE_CONFIG에서 첫 번째 접미사의 설정을 자동 로드
        if RANGE_CONFIG:
            first_suffix = sorted(RANGE_CONFIG.keys())[0]
            self.load_config_for_suffix(first_suffix)
        else:
            # RANGE_CONFIG가 비어있으면 기본값 설정
            self.mapping_table.setRowCount(1)
            self.mapping_table.setItem(0, 0, QTableWidgetItem('GSM 850'))
            self.mapping_table.setItem(0, 1, QTableWidgetItem('B27:I31'))
            self.mapping_table.setItem(0, 2, QTableWidgetItem('GSM 850'))

        mapping_layout.addWidget(self.mapping_table)

        # 버튼
        btn_layout = QHBoxLayout()
        add_btn = QPushButton("+ 행 추가")
        add_btn.clicked.connect(self.add_mapping_row)
        btn_layout.addWidget(add_btn)

        del_btn = QPushButton("- 행 삭제")
        del_btn.clicked.connect(self.delete_mapping_row)
        btn_layout.addWidget(del_btn)

        btn_layout.addStretch()
        mapping_layout.addLayout(btn_layout)

        left_column.addWidget(mapping_group)

        # 좌측 컬럼을 content_layout에 추가
        content_layout.addLayout(left_column, 60)  # 60% 너비

        # 우측 컬럼 (로그)
        right_column = QVBoxLayout()
        right_column.setSpacing(10)

        # 사용 가이드
        guide_group = QGroupBox("📖 사용 가이드")
        guide_layout = QVBoxLayout()

        guide_text = QTextEdit()
        guide_text.setReadOnly(True)
        guide_text.setMaximumHeight(200)
        guide_text.setStyleSheet("""
            QTextEdit {
                background-color: #2c3e50;
                color: #ecf0f1;
                border: 1px solid #34495e;
                border-radius: 4px;
                padding: 10px;
                font-size: 9pt;
                line-height: 1.6;
            }
        """)
        guide_text.setHtml("""
<h3 style='color: #3498db;'>📖 Excel 범위 삽입 사용 방법</h3>

<p><b style='color: #27ae60;'>【1. 설정 관리】</b></p>
<p style='margin-left: 15px;'>
• 설정 파일에서 범위 설정 불러오기<br>
• 테이블 편집 후 저장 가능<br>
• 접미사별로 다른 설정 관리 (#1, #2 등)
</p>

<p><b style='color: #3498db;'>【2. 파일 선택】</b></p>
<p style='margin-left: 15px;'>
• Excel 파일: 데이터가 있는 파일 (다중 선택)<br>
• Word 파일: 마커가 있는 템플릿 (다중 선택)<br>
• 파일명 형식: 파일명_#1.xlsx
</p>

<p><b style='color: #e67e22;'>【3. 실행】</b></p>
<p style='margin-left: 15px;'>
• "실행" 버튼 클릭<br>
• Excel 범위가 Word 마커에 자동 삽입<br>
• 복사본 파일 자동 생성 (_copy)
</p>
        """)
        guide_layout.addWidget(guide_text)
        guide_group.setLayout(guide_layout)
        right_column.addWidget(guide_group)

        # 로그
        log_group = QGroupBox("📝 처리 로그")
        log_layout = QVBoxLayout()
        log_layout.setSpacing(8)
        log_group.setLayout(log_layout)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setFont(QFont("Consolas", 9))
        self.log_text.setStyleSheet("""
            QTextEdit {
                border: 1px solid #dcdde1;
                border-radius: 4px;
                background-color: #2c3e50;
                color: #2ecc71;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 9pt;
                padding: 8px;
            }
        """)
        log_layout.addWidget(self.log_text)

        clear_btn = QPushButton("로그 지우기")
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        clear_btn.clicked.connect(self.log_text.clear)
        log_layout.addWidget(clear_btn)

        right_column.addWidget(log_group)

        # 우측 컬럼을 content_layout에 추가
        content_layout.addLayout(right_column, 40)  # 40% 너비

        # content_layout을 main_layout에 추가
        main_layout.addLayout(content_layout)

        # 실행 버튼
        self.run_btn = QPushButton("▶ 실행")
        self.run_btn.setMinimumHeight(45)
        self.run_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px;
                border-radius: 6px;
                font-weight: bold;
                font-size: 12pt;
            }
            QPushButton:hover {
                background-color: #229954;
            }
            QPushButton:disabled {
                background-color: #bdc3c7;
            }
        """)
        self.run_btn.clicked.connect(self.run_process)
        main_layout.addWidget(self.run_btn)

        # 진행 표시줄
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setRange(0, 0)
        self.progress_bar.setMinimumHeight(25)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #3498db;
                border-radius: 5px;
                text-align: center;
                background-color: #ecf0f1;
                font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #27ae60;
                border-radius: 4px;
            }
        """)
        main_layout.addWidget(self.progress_bar)

    # ========== FILE SELECTION METHODS ==========

    def select_excel_files(self):
        """엑셀 파일 다중 선택"""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "엑셀 파일 선택 (다중 선택 가능)",
            "",
            "Excel Files (*.xlsx *.xls)"
        )
        if file_paths:
            self.excel_files = sorted(file_paths)  # 정렬
            # 파일명만 추출해서 표시
            file_names = [os.path.basename(f) for f in self.excel_files]
            self.excel_edit.setText(", ".join(file_names))
            self.selected_files_label.setText(f"선택된 파일: {len(self.excel_files)}개")
            self.log_text.append(f"\n✓ {len(self.excel_files)}개 파일 선택됨:")
            for i, file_name in enumerate(file_names, 1):
                self.log_text.append(f"  {i}. {file_name}")

    def select_word_files(self):
        """워드 파일 다중 선택"""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "워드 파일 선택 (다중 선택 가능)",
            "",
            "Word Files (*.docx)"
        )
        if file_paths:
            self.word_files = sorted(file_paths)  # 정렬
            # 파일명만 추출해서 표시
            file_names = [os.path.basename(f) for f in self.word_files]
            self.word_edit.setText(", ".join(file_names))
            self.selected_word_files_label.setText(f"선택된 파일: {len(self.word_files)}개")
            self.log_text.append(f"\n✓ {len(self.word_files)}개 워드 파일 선택됨:")
            for i, file_name in enumerate(file_names, 1):
                self.log_text.append(f"  {i}. {file_name}")

    # ========== TABLE MANAGEMENT METHODS ==========

    def add_mapping_row(self):
        """매핑 행 추가"""
        row = self.mapping_table.rowCount()
        self.mapping_table.insertRow(row)

    def delete_mapping_row(self):
        """선택된 행 삭제"""
        current_row = self.mapping_table.currentRow()
        if current_row >= 0:
            self.mapping_table.removeRow(current_row)

    def load_config_to_table(self):
        """선택된 접미사의 RANGE_CONFIG를 테이블에 로드"""
        suffix = self.suffix_combo.currentText()
        self.load_config_for_suffix(suffix)

    def load_config_for_suffix(self, suffix):
        """특정 접미사의 설정을 테이블에 로드"""
        if suffix in RANGE_CONFIG:
            config_list = RANGE_CONFIG[suffix]
            self.mapping_table.setRowCount(len(config_list))

            for row, config in enumerate(config_list):
                self.mapping_table.setItem(row, 0, QTableWidgetItem(config['sheet']))
                self.mapping_table.setItem(row, 1, QTableWidgetItem(config['range']))
                self.mapping_table.setItem(row, 2, QTableWidgetItem(config['marker']))
                self.mapping_table.setItem(row, 3, QTableWidgetItem(config.get('category', '')))

            # log_text가 존재하는 경우에만 로그 출력
            if hasattr(self, 'log_text'):
                self.log_text.append(f"✓ {suffix} 설정 로드됨: {len(config_list)}개 항목")
        else:
            if hasattr(self, 'log_text'):
                self.log_text.append(f"⚠️ {suffix} 설정이 없습니다.")

    # ========== CONFIGURATION MANAGEMENT METHODS ==========

    def reload_config_file(self):
        """설정 파일 새로고침"""
        global RANGE_CONFIG

        if not os.path.exists(CONFIG_FILE_PATH):
            QMessageBox.warning(self, "경고", f"설정 파일을 찾을 수 없습니다:\n{CONFIG_FILE_PATH}")
            return

        loaded_config = self.load_config_from_excel()
        if loaded_config:
            RANGE_CONFIG = loaded_config
            # 콤보박스 업데이트
            self.suffix_combo.clear()
            self.suffix_combo.addItems(sorted(RANGE_CONFIG.keys()))
            self.log_text.append(f"✓ 설정 파일 새로고침 완료: {len(RANGE_CONFIG)}개 접미사")
            QMessageBox.information(self, "완료", "설정 파일을 새로고침했습니다.")
        else:
            QMessageBox.critical(self, "오류", "설정 파일 로드에 실패했습니다.")

    def save_current_config(self):
        """현재 테이블을 설정 파일로 저장"""
        global RANGE_CONFIG

        # 현재 선택된 접미사
        current_suffix = self.suffix_combo.currentText()

        if not current_suffix:
            QMessageBox.warning(self, "경고", "접미사를 선택하세요.")
            return

        # 테이블에서 설정 가져오기
        config_list = self.get_mappings()
        if not config_list:
            QMessageBox.warning(self, "경고", "테이블이 비어있습니다.")
            return

        # RANGE_CONFIG 업데이트
        RANGE_CONFIG[current_suffix] = config_list

        # 엑셀 파일로 저장
        if self.save_config_to_excel_file(RANGE_CONFIG, CONFIG_FILE_PATH):
            self.log_text.append(f"✓ {current_suffix} 설정 저장 완료: {len(config_list)}개 항목")
            QMessageBox.information(self, "완료", f"설정을 파일로 저장했습니다:\n{CONFIG_FILE_PATH}")
        else:
            QMessageBox.critical(self, "오류", "설정 파일 저장에 실패했습니다.")

    def open_config_file(self):
        """설정 파일을 엑셀로 열기"""
        if not os.path.exists(CONFIG_FILE_PATH):
            QMessageBox.warning(self, "경고", f"설정 파일을 찾을 수 없습니다:\n{CONFIG_FILE_PATH}")
            return

        try:
            os.startfile(CONFIG_FILE_PATH)
            self.log_text.append(f"✓ 설정 파일 열기: {CONFIG_FILE_NAME}")
        except Exception as e:
            QMessageBox.critical(self, "오류", f"파일 열기 실패:\n{str(e)}")

    def get_mappings(self):
        """테이블에서 매핑 정보 가져오기"""
        mappings = []
        for row in range(self.mapping_table.rowCount()):
            sheet = self.mapping_table.item(row, 0)
            range_addr = self.mapping_table.item(row, 1)
            marker = self.mapping_table.item(row, 2)
            category = self.mapping_table.item(row, 3)

            if sheet and range_addr and marker:
                sheet_text = sheet.text().strip()
                range_text = range_addr.text().strip()
                marker_text = marker.text().strip()
                category_text = category.text().strip() if category else ''

                if sheet_text and range_text and marker_text:
                    item = {
                        'sheet': sheet_text,
                        'range': range_text,
                        'marker': marker_text
                    }
                    if category_text:
                        item['category'] = category_text
                    mappings.append(item)

        return mappings

    # ========== PROCESS EXECUTION ==========

    def run_process(self):
        """프로세스 실행"""
        # 엑셀 파일 확인
        if not self.excel_files:
            QMessageBox.warning(self, "경고", "엑셀 데이터 파일을 선택하세요.")
            return

        # 워드 파일 확인
        if not self.word_files:
            QMessageBox.warning(self, "경고", "워드 템플릿 파일을 선택하세요.")
            return

        # 파일 존재 확인
        missing_files = [f for f in self.excel_files if not os.path.exists(f)]
        if missing_files:
            QMessageBox.warning(self, "경고", f"다음 엑셀 파일을 찾을 수 없습니다:\n{', '.join([os.path.basename(f) for f in missing_files])}")
            return

        missing_word_files = [f for f in self.word_files if not os.path.exists(f)]
        if missing_word_files:
            QMessageBox.warning(self, "경고", f"다음 워드 파일을 찾을 수 없습니다:\n{', '.join([os.path.basename(f) for f in missing_word_files])}")
            return

        mappings = self.get_mappings()
        # mappings가 비어있어도 RANGE_CONFIG 사용 가능하므로 경고만 표시
        if not mappings:
            self.log_text.append("⚠️ GUI 테이블이 비어있습니다. RANGE_CONFIG 설정을 사용합니다.")

        # UI 업데이트
        self.run_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.log_text.clear()

        # 상태 바 업데이트
        self.status_update.emit("⏳ 처리 중... 잠시만 기다려주세요")

        # 워커 스레드 시작
        self.worker = ExcelRangeProcessorThread(self.excel_files, self.word_files, mappings)
        self.worker.progress.connect(self.update_log)
        self.worker.finished.connect(self.process_finished)
        self.worker.start()

    def update_log(self, message):
        """로그 업데이트"""
        self.log_text.append(message)
        self.log_text.verticalScrollBar().setValue(
            self.log_text.verticalScrollBar().maximum()
        )

    def process_finished(self, result):
        """처리 완료"""
        self.run_btn.setEnabled(True)
        self.progress_bar.setVisible(False)

        if result['success']:
            # 처리 시간 포맷팅
            elapsed_time = result.get('elapsed_time', 0)
            minutes = int(elapsed_time // 60)
            seconds = int(elapsed_time % 60)

            # 상태 바 업데이트
            self.status_update.emit(f"✅ 완료 - 처리 시간: {minutes}분 {seconds}초")

            msg = (
                f"처리 완료!\n\n"
                f"⏱️ 처리 시간: {minutes}분 {seconds}초\n"
                f"생성된 워드 파일: {len(result['output_files'])}개\n"
                f"삽입 성공: {result['images_inserted']}개\n"
                f"삽입 실패: {result['images_failed']}개\n"
            )

            # 실패한 마커가 있으면 추가 정보 표시
            if result['failed_markers']:
                failed_list = "\n".join([f"  • {m['marker']}" for m in result['failed_markers'][:5]])
                if len(result['failed_markers']) > 5:
                    failed_list += f"\n  ... 외 {len(result['failed_markers']) - 5}개"
                msg += f"\n실패한 마커:\n{failed_list}\n\n(자세한 내용은 로그 확인)\n"

            if result['output_files']:
                msg += f"\n생성된 파일 목록:\n"
                for f in result['output_files'][:3]:
                    msg += f"  • {os.path.basename(f)}\n"
                if len(result['output_files']) > 3:
                    msg += f"  ... 외 {len(result['output_files']) - 3}개"

            QMessageBox.information(self, "완료", msg)
        else:
            # 상태 바 업데이트
            self.status_update.emit("❌ 처리 실패")
            QMessageBox.critical(self, "오류", result['message'])

    # ========== HELPER METHODS ==========

    def load_or_create_config(self):
        """설정 파일 로드 또는 생성"""
        if os.path.exists(CONFIG_FILE_PATH):
            return self.load_config_from_excel()
        else:
            self.create_default_config_file()
            return DEFAULT_RANGE_CONFIG.copy()

    def load_config_from_excel(self):
        """Excel 설정 파일 로드"""
        try:
            wb = load_workbook(CONFIG_FILE_PATH)
            config = {}

            for sheet_name in wb.sheetnames:
                if sheet_name.startswith("#"):
                    suffix = sheet_name
                    config[suffix] = []

                    ws = wb[sheet_name]
                    for row in ws.iter_rows(min_row=2, values_only=True):
                        if row[0]:  # Sheet name exists
                            config[suffix].append({
                                "sheet": str(row[0]),
                                "range": str(row[1]),
                                "marker": str(row[2]),
                                "category": str(row[3]) if len(row) > 3 and row[3] else ""
                            })

            wb.close()
            logger.info(f"✓ 설정 파일 로드 완료: {len(config)}개 접미사")
            return config
        except Exception as e:
            logger.error(f"설정 파일 로드 실패: {e}")
            return DEFAULT_RANGE_CONFIG.copy()

    def create_default_config_file(self):
        """기본 설정 파일 생성"""
        try:
            from openpyxl.styles import Font, PatternFill, Alignment

            wb = Workbook()
            wb.remove(wb.active)

            for suffix, items in DEFAULT_RANGE_CONFIG.items():
                ws = wb.create_sheet(title=suffix)
                ws.append(["Sheet Name", "Range", "Marker", "Category"])

                # 헤더 스타일
                header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                header_font = Font(bold=True, color="FFFFFF")
                for cell in ws[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')

                for item in items:
                    ws.append([
                        item["sheet"],
                        item["range"],
                        item["marker"],
                        item.get("category", "")
                    ])

                # 열 너비 자동 조정
                ws.column_dimensions['A'].width = 20
                ws.column_dimensions['B'].width = 15
                ws.column_dimensions['C'].width = 25
                ws.column_dimensions['D'].width = 30

            wb.save(CONFIG_FILE_PATH)
            logger.info(f"기본 설정 파일 생성: {CONFIG_FILE_PATH}")
        except Exception as e:
            logger.error(f"설정 파일 생성 실패: {e}")

    def save_config_to_excel_file(self, config, file_path):
        """RANGE_CONFIG를 엑셀 파일로 저장"""
        try:
            from openpyxl.styles import Font, PatternFill, Alignment

            wb = Workbook()
            wb.remove(wb.active)  # 기본 시트 제거

            for suffix, config_list in sorted(config.items()):
                ws = wb.create_sheet(title=suffix)

                # 헤더 작성
                ws['A1'] = 'Sheet Name'
                ws['B1'] = 'Range'
                ws['C1'] = 'Marker'
                ws['D1'] = 'Category'

                # 헤더 스타일
                header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                header_font = Font(bold=True, color="FFFFFF")
                for cell in ws[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')

                # 카테고리별로 그룹화
                categories = {}
                for item in config_list:
                    category = item.get('category', '기타')
                    if category not in categories:
                        categories[category] = []
                    categories[category].append(item)

                # 데이터 작성 (카테고리별 구분)
                current_row = 2
                category_fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
                category_font = Font(bold=True, size=11)

                for category, items in categories.items():
                    # 카테고리 행 삽입
                    ws[f'A{current_row}'] = f"【{category}】"
                    ws.merge_cells(f'A{current_row}:D{current_row}')
                    category_cell = ws[f'A{current_row}']
                    category_cell.font = category_font
                    category_cell.fill = category_fill
                    category_cell.alignment = Alignment(horizontal='left')
                    current_row += 1

                    # 카테고리 항목 작성
                    for item in items:
                        ws[f'A{current_row}'] = item['sheet']
                        ws[f'B{current_row}'] = item['range']
                        ws[f'C{current_row}'] = item['marker']
                        ws[f'D{current_row}'] = item.get('category', '')
                        current_row += 1

                    # 카테고리 사이 빈 행
                    current_row += 1

                # 열 너비 자동 조정
                ws.column_dimensions['A'].width = 20
                ws.column_dimensions['B'].width = 15
                ws.column_dimensions['C'].width = 25
                ws.column_dimensions['D'].width = 30

            wb.save(file_path)
            wb.close()
            logger.info(f"✓ 설정 파일 저장 완료: {file_path}")
            return True

        except Exception as e:
            logger.error(f"✗ 설정 파일 저장 실패: {str(e)}")
            return False


# ===================================================================
# MAIN ENTRY POINT
# ===================================================================

def main():
    """메인 애플리케이션 실행"""
    app = QApplication(sys.argv)

    # Set application font
    font = QFont("Malgun Gothic", 9)
    app.setFont(font)

    # Create and show main window
    window = IntegratedWordExcelManager()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
